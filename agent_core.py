#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
智能体模块 (agents.py)
=====================
定义招聘场景中的两个核心 AI 智能体

包含类：
1. MatcherAgent - 简历与职位匹配智能体
2. InterviewerAgent - 模拟面试智能体

作者：AI助手
"""
from dotenv import load_dotenv
load_dotenv()
import time  # 用于添加延迟，让对话更自然
import re    # 用于正则表达式解析
import json  # 用于JSON格式解析
import os     # 用于环境变量
from datetime import datetime  # 用于时间戳
from typing import Dict, Any, List
from vector_db import (
    load_jobs,           # 加载职位数据
    init_vector_db,      # 初始化向量数据库
    format_job_text,     # 格式化职位文本
    format_resume_for_matching  # 格式化简历文本用于匹配
)
from resume_parser import ResumeParser  # 简历解析器

# RAG 知识库（延迟导入，避免循环依赖和初始化顺序问题）
_rag_chain = None


def _get_rag():
    """延迟获取 RAGChain 实例"""
    global _rag_chain
    if _rag_chain is None:
        try:
            from knowledge_base import get_rag_chain
            _rag_chain = get_rag_chain()
        except Exception as e:
            print(f"[RAG] Knowledge base init skipped: {e}")
            _rag_chain = False  # 标记为不可用
    return _rag_chain if _rag_chain is not False else None


# =============================================================================
# 第一部分：LLM 调用函数
# =============================================================================

def ask_llm(prompt: str, model: str = "deepseek-v4-pro") -> str:
    print(f"DEBUG: API Key = {os.getenv('DEEPSEEK_API_KEY')[:10]}...")
    """
    调用阿里云百炼 LLM API（OpenAI 兼容模式）

    参数：
        prompt: 提示词
        model: 模型名称

    返回：
        str: API 返回的回答
    """
    from openai import OpenAI

    api_key = os.getenv("DEEPSEEK_API_KEY")
    if not api_key:
        raise ValueError("未设置 DEEPSEEK_API_KEY 环境变量")

    client = OpenAI(
        api_key=api_key,
        base_url="https://api.deepseek.com/v1",
        timeout=120.0,
    )

    try:
        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
        )
        return response.choices[0].message.content
    except Exception as e:
        raise RuntimeError(f"LLM API 调用失败: {str(e)}")


# 保留旧名称作为别名，兼容现有调用
def ask_glm(prompt: str, model: str = "deepseek-v4-pro") -> str:
    """兼容旧接口，内部调用 ask_llm"""
    return ask_llm(prompt, model=model)


# =============================================================================
# 第二部分：MatcherAgent - 简历与职位匹配智能体
# =============================================================================

class MatcherAgent:
    """
    简历与职位匹配智能体
    
    职责：
    1. 加载和初始化向量数据库
    2. 计算简历与职位的匹配度
    3. 生成专业的推荐理由
    
    工作流程：
    输入简历 → 向量检索 → 计算相似度 → LLM评价 → 输出结果
    """
    
    def __init__(self):
        """
        初始化匹配器
        """
        self._vector_db = None  # 向量数据库
        self._jobs_data = None  # 职位数据
        self._initialized = False  # 是否已初始化
    
    def initialize(self, jobs_file='jobs.json', force_reset=True):
        """
        初始化智能体
        
        参数：
            jobs_file: 职位数据文件路径
            force_reset: 是否强制重建数据库
        
        示例：
            >>> agent = MatcherAgent()
            >>> agent.initialize('jobs.json')
            正在加载职位数据...
            已加载 3 个职位
        """
        if self._initialized:
            print("数据库已初始化，跳过重复初始化")
            return
        
        print("正在加载职位数据...")
        self._jobs_data = load_jobs(jobs_file)
        print(f"已加载 {len(self._jobs_data)} 个职位")
        
        print("\n正在初始化向量数据库...")
        self._vector_db, _ = init_vector_db(jobs_file, force_reset=force_reset)
        print("向量数据库初始化完成！")
        
        self._initialized = True
    
    def _calculate_similarity_score(self, distance, max_distance=30.0, baseline=0.0):
        """
        计算相似度分数
        
        参数：
            distance: 向量距离（越小越相似）
            max_distance: 最大可能距离
            baseline: 基线值（低于此值的相似度归零）
        
        返回：
            float: 0-100 的匹配分数
        """
        # 将距离转换为相似度
        similarity = max(0, min(1, (max_distance - distance) / max_distance))
        
        # 应用基线
        if similarity < baseline:
            return 0.0
        
        # 转换为百分比分数
        score = ((similarity - baseline) / (1.0 - baseline)) * 100
        return max(0, min(100, round(score, 2)))
    
    def match_resume(self, resume, top_k=3):
        """
        匹配单个简历与所有职位
        
        参数：
            resume: 简历字典
            top_k: 返回前几名匹配结果（默认前3）
        
        返回：
            list: 匹配结果列表，按匹配度降序排列
        
        示例：
            >>> agent.match_resume(resume)
            [
                {
                    'match_score': 85.5,
                    'job_title': 'Python后端开发工程师',
                    'company': '创新科技有限公司',
                    'reason': '技能匹配度高...'
                }
            ]
        """
        if not self._initialized:
            raise RuntimeError("MatcherAgent 尚未初始化，请先调用 initialize()")
        
        # 格式化简历文本用于匹配
        resume_text = format_resume_for_matching(resume)
        
        # 在向量数据库中搜索相似职位
        # k=top_k * 2 因为每个职位存了两份（完整描述+技能描述）
        results = self._vector_db.similarity_search_with_score(
            resume_text, 
            k=top_k * 2
        )
        
        if not results:
            return []
        
        # 按职位ID分组
        job_scores = {}
        for doc, distance in results:
            # 安全获取元数据，使用.get()防止KeyError
            job_id = doc.metadata.get('job_id', 'unknown')
            
            if job_id not in job_scores:
                job_scores[job_id] = {
                    'job_id': job_id,
                    'job_title': doc.metadata.get('title', '未知职位'),
                    'company': doc.metadata.get('company', '未知公司'),
                    'location': doc.metadata.get('location', '未知地点'),
                    'job_content': doc.page_content,
                    'sources': [],
                    'distances': []
                }
            
            job_scores[job_id]['sources'].append(doc.metadata.get('source', 'full_description'))
            job_scores[job_id]['distances'].append(distance)
        
        # 合并并计算分数
        merged_results = []
        for job_id, data in job_scores.items():
            # 计算平均距离
            avg_distance = sum(data['distances']) / len(data['distances'])
            
            # 计算匹配分数
            match_score = self._calculate_similarity_score(avg_distance)
            
            # 找到对应的职位详细信息（安全比较，处理空格问题）
            job_info = None
            job_id_str = str(job_id).strip()  # 去除首尾空格
            
            for j in self._jobs_data:
                # 安全获取ID并去除空格
                j_id = str(j.get('id', '')).strip()
                if j_id == job_id_str:
                    job_info = j
                    break
            
            if job_info:
                # 生成推荐理由（传入匹配分数）
                reason_data = self._generate_reason(resume, job_info, match_score)
                
                # 将 reason 对象转换为字符串格式（适配前端期望）
                reason_str = f"【匹配度】{reason_data.get('match_level', '一般')}\n"
                reason_str += f"【技能匹配】{reason_data.get('skill_match', '')}\n"
                reason_str += f"【经验匹配】{reason_data.get('experience_match', '')}\n"
                reason_str += f"【综合评价】{reason_data.get('overall_evaluation', '')}\n"
                reason_str += f"【建议】{reason_data.get('suggestion', '')}"
                
                merged_results.append({
                    'id': job_id,  # 前端期望的字段名
                    'title': data['job_title'],  # 前端期望的字段名
                    'company': data['company'],
                    'location': data['location'],
                    'match_score': match_score,
                    'reason': reason_str,  # 转换为字符串
                    'responsibilities': job_info.get('responsibilities', ''),
                    'requirements': job_info.get('requirements', ''),
                    'job_info': job_info  # 保留完整职位信息供面试使用
                })
        
        # 按匹配分数降序排序
        merged_results.sort(key=lambda x: x['match_score'], reverse=True)
        
        return merged_results[:top_k]
    
    def _generate_reason(self, resume, job, match_score=None):
        """
        使用 LLM 生成推荐理由（返回结构化 JSON）
        
        参数：
            resume: 简历字典
            job: 职位字典
            match_score: 匹配分数（可选）
        
        返回：
            dict: 结构化的推荐理由字典
        """
        resume_text = format_resume_for_matching(resume)
        job_text = format_job_text(job)
        
        prompt = f"""请分析以下简历与职位的匹配度，并返回结构化的分析报告：

【简历】
{resume_text}

【职位】
{job_text}

【匹配分数】{match_score if match_score is not None else '未知'}

请返回JSON格式，包含以下字段：
{{
    "skill_match": "技能匹配分析（1-3句话）",
    "experience_match": "经验匹配分析（1-3句话）",
    "overall_evaluation": "综合评价（1-2句话）",
    "suggestion": "面试建议（1-2句话）",
    "match_level": "匹配等级：优秀/良好/一般/较差"
}}

注意：只输出JSON，不要有其他文字。
"""
        
        try:
            response = ask_glm(prompt)
            # 尝试解析JSON
            result = self._parse_json_response(response)
            if result:
                return result
            # 如果JSON解析失败，返回简单文本格式
            return {
                'skill_match': response[:100] + '...' if len(response) > 100 else response,
                'experience_match': '',
                'overall_evaluation': '',
                'suggestion': '',
                'match_level': '一般'
            }
        except Exception as e:
            print(f"警告：LLM调用失败: {str(e)}")
            return {
                'skill_match': '技能匹配分析中',
                'experience_match': '经验匹配分析中',
                'overall_evaluation': '基于简历与职位的分析',
                'suggestion': '建议深入了解职位需求',
                'match_level': '一般'
            }
    
    def _parse_json_response(self, response):
        """
        解析LLM返回的JSON响应
        
        参数：
            response: LLM返回的字符串
        
        返回：
            dict: 解析后的字典，如果解析失败返回None
        """
        try:
            # 尝试提取JSON部分（可能在markdown代码块中）
            if '```json' in response:
                json_str = response.split('```json')[1].split('```')[0].strip()
            elif '```' in response:
                json_str = response.split('```')[1].split('```')[0].strip()
            else:
                json_str = response.strip()
            
            # 尝试解析JSON
            return json.loads(json_str)
        except json.JSONDecodeError:
            # JSON解析失败，尝试提取引号内的内容作为备选
            try:
                # 提取所有双引号内的内容
                quoted_strings = re.findall(r'"([^"]+)"', response)
                if len(quoted_strings) >= 3:
                    return {
                        'skill_match': quoted_strings[0] if len(quoted_strings) > 0 else '',
                        'experience_match': quoted_strings[1] if len(quoted_strings) > 1 else '',
                        'overall_evaluation': quoted_strings[2] if len(quoted_strings) > 2 else '',
                        'suggestion': quoted_strings[3] if len(quoted_strings) > 3 else '',
                        'match_level': quoted_strings[4] if len(quoted_strings) > 4 else '一般'
                    }
            except Exception as e:
                print(f"警告：JSON解析失败，备选方案也失败: {str(e)}")
            return None
        except Exception as e:
            print(f"警告：解析响应失败: {str(e)}")
            return None
    
    def match_all_resumes(self, resumes_file):
        """
        批量匹配所有简历

        参数：
            resumes_file: 简历文件路径（支持 JSON、PDF、DOCX 格式）

        返回：
            list: 所有简历的匹配结果
        """
        if not self._initialized:
            raise RuntimeError("MatcherAgent 尚未初始化，请先调用 initialize()")

        # 根据文件扩展名选择加载方式
        ext = os.path.splitext(resumes_file)[1].lower()

        if ext in ['.pdf', '.docx', '.doc']:
            # 使用 resume_parser 解析文件
            try:
                from resume_parser import ResumeParser
                parser = ResumeParser()
                parsed = parser.parse_file(resumes_file)
                resumes = [parsed]
            except Exception as e:
                print(f"警告：解析简历文件失败: {str(e)}")
                return []
        elif ext == '.json':
            # 从 JSON 文件加载
            try:
                with open(resumes_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                resumes = data.get('resumes', [])
            except Exception as e:
                print(f"警告：加载简历 JSON 文件失败: {str(e)}")
                return []
        else:
            print(f"警告：不支持的文件格式: {ext}")
            return []

        if not resumes:
            print("警告：未加载到任何简历数据")
            return []

        results = []
        total = len(resumes)

        print(f"\n正在匹配 {total} 份简历...")

        for idx, resume in enumerate(resumes, 1):
            try:
                print(f"\r匹配进度: {idx}/{total}", end='', flush=True)

                match_results = self.match_resume(resume)

                results.append({
                    'resume': resume,
                    'matches': match_results
                })

                # 添加小延迟，避免请求过快
                time.sleep(0.5)

            except Exception as e:
                print(f"\n警告：处理简历 {resume.get('name', f'简历{idx}')} 时发生错误: {str(e)}")

        print("\n匹配完成！")
        return results


# =============================================================================
# 第二部分：InterviewerAgent - 模拟面试智能体
# =============================================================================

class InterviewerAgent:
    """
    模拟面试智能体
    
    职责：
    1. 基于简历和目标职位生成面试问题
    2. 接收候选人的回答
    3. 进行评价和追问
    4. 生成面试总结
    
    面试流程：
    开始面试 → 生成问题 → 候选人回答 → AI评价 → 追问（可选）→ 面试总结
    """
    
    def __init__(self):
        """
        初始化面试官智能体
        """
        self.current_resume = None    # 当前面试的简历
        self.current_job = None       # 当前面试的目标职位
        self.questions = []           # 生成的面试问题
        self.answers = []             # 候选人的回答
        self.current_question_idx = 0 # 当前问题索引
    
    def start_interview(self, resume, job, memory_context=None,match_score=70):
        """
        开始一场模拟面试
        
        参数：
            resume: 简历字典
            job: 职位字典（包括 job_info 字段）
            memory_context: 历史记忆上下文（可选）
        
        示例：
            >>> agent = InterviewerAgent()
            >>> agent.start_interview(resume, matched_job)
        """
        # 安全获取简历信息，使用默认值防止KeyError
        self.current_resume = {
            'id': resume.get('id', 'unknown'),
            'name': resume.get('name', '未知候选人'),
            'skills': resume.get('skills', ''),
            'experience': resume.get('experience', '')
        }
        
        # 存储记忆上下文
        self.memory_context = memory_context
        
        # 安全获取职位信息，优先使用 job_info
        job_info = job.get('job_info') or job
        self.current_job = {
            'id': job_info.get('id', 'unknown'),
            'title': job_info.get('title', '未知职位'),
            'company': job_info.get('company', '未知公司'),
            'location': job_info.get('location', '未知地点'),
            'responsibilities': job_info.get('responsibilities', ''),
            'requirements': job_info.get('requirements', '')
        }
        
        self.questions = []
        self.answers = []
        self.current_question_idx = 0
        self._interview_ended = False  # 新增：标记面试是否已结束
        self.match_score = match_score  # ✅ 保存匹配分数为实例属性
        
        print("\n" + "="*60)
        print(f"🤖 模拟面试开始")
        print(f"👤 候选人：{self.current_resume['name']}")
        print(f"💼 目标职位：{self.current_job['title']} @ {self.current_job['company']}")
        print("="*60)
        
        # 生成面试问题
        self._generate_questions()
        
        # 显示第一个问题
        self._display_current_question()
    
    def _generate_questions(self):
        """
        使用 LLM 生成面试问题
        
        优先使用JSON格式解析，失败则使用正则表达式解析
        根据匹配分数动态调整问题难度
        """
        resume_text = format_resume_for_matching(self.current_resume)
        job_text = format_job_text(self.current_job)
        
        # 获取匹配分数（用于动态调整问题难度）
        match_score = self.match_score
        difficulty_instruction = ""
        
        # 优化：要求LLM输出JSON格式，便于解析
        # 根据匹配分数调整问题难度
        difficulty_instruction = ""
        if match_score is not None:
            if match_score >= 80:
                difficulty_instruction = f"【难度调整】候选人匹配度很高（得分{match_score}），请侧重深度问题，考察技术细节和复杂场景处理能力。"
            elif match_score >= 60:
                difficulty_instruction = f"【难度调整】候选人匹配度良好（得分{match_score}），请平衡广度和深度问题。"
            else:
                difficulty_instruction = f"【难度调整】候选人匹配度一般（得分{match_score}），请侧重基础和广度问题，考察基础知识和学习能力。"
        
        # 构建记忆上下文提示
        memory_prompt = ""
        if getattr(self, 'memory_context', None):
            memory_prompt = f"""【历史记录】
{self.memory_context}

请根据历史反馈设计问题，特别是针对之前发现的弱点进行考察。
"""

        # ====== RAG 增强：检索面试知识库 ======
        rag_context = ""
        try:
            rag = _get_rag()
            if rag:
                resume_skills = self.current_resume.get('skills', [])
                if isinstance(resume_skills, str):
                    resume_skills = [s.strip() for s in resume_skills.split(',')]
                job_title = self.current_job.get('title', '')
                rag_context = rag.retrieve_for_question_generation(
                    job_title=job_title,
                    resume_skills=resume_skills,
                    top_k=4,
                )
                if rag_context:
                    print(f"   [RAG] Retrieved {rag_context.count('### 参考')} knowledge chunks for question generation")
        except Exception as e:
            print(f"   [RAG] Knowledge retrieval skipped: {e}")
            rag_context = ""
        # ====== RAG 增强结束 ======

        prompt = f"""作为一位资深面试官，请为以下候选人针对该职位生成5个面试问题。

{rag_context}

【候选人简历】
{resume_text}

【目标职位】
{job_text}

{difficulty_instruction}

{memory_prompt}

请生成5个有针对性的面试问题，涵盖：
1. 技术技能相关问题（基于简历技能和职位要求）
2. 项目经验相关问题（基于简历中的工作经历）
3. 问题解决能力相关问题
4. 团队协作相关问题
5. 职业规划相关问题

请以JSON格式输出，格式如下：
{{
    "questions": [
        "问题1的内容",
        "问题2的内容",
        "问题3的内容",
        "问题4的内容",
        "问题5的内容"
    ]
}}

注意：只输出JSON，不要有其他文字。
"""
        
        try:
            response = ask_glm(prompt)
            
            # 优先尝试JSON格式解析
            questions = self._parse_json_questions(response)
            
            if questions:
                self.questions = questions
            else:
                # JSON解析失败，使用正则表达式解析
                self.questions = self._parse_regex_questions(response)
            
            # 如果仍然解析失败，使用默认问题
            if not self.questions:
                self.questions = self._get_default_questions()
                
        except Exception as e:
            print(f"警告：生成问题失败，使用默认问题: {str(e)}")
            self.questions = self._get_default_questions()
    
    def _parse_json_questions(self, response):
        """
        尝试解析JSON格式的问题列表
        
        参数：
            response: LLM返回的响应文本
        
        返回：
            list: 问题列表，如果解析失败返回None
        """
        try:
            # 尝试提取JSON部分
            json_str = response.strip()
            
            # 移除可能存在的markdown代码块标记
            if json_str.startswith('```'):
                # 找到第一个{的位置
                start = json_str.find('{')
                end = json_str.rfind('}')
                if start != -1 and end != -1:
                    json_str = json_str[start:end+1]
            
            # 解析JSON
            data = json.loads(json_str)
            
            if isinstance(data, dict) and 'questions' in data:
                questions = data['questions']
                if isinstance(questions, list) and len(questions) >= 3:
                    # 验证每个问题都是非空字符串
                    valid_questions = [q.strip() for q in questions if isinstance(q, str) and len(q.strip()) > 5]
                    if len(valid_questions) >= 3:
                        return valid_questions[:5]  # 最多返回5个问题
            
            # 如果questions字段存在但格式不对，尝试提取引号内内容
            return self._extract_quoted_questions(json_str)
            
        except (json.JSONDecodeError, Exception) as e:
            print(f"JSON解析失败: {str(e)}")
            # JSON解析失败，尝试提取引号内的内容作为备选方案
            return self._extract_quoted_questions(response)
    
    def _extract_quoted_questions(self, text):
        """
        从文本中提取所有引号内的内容作为问题备选方案
        
        参数：
            text: 输入文本
        
        返回：
            list: 提取的问题列表，如果失败返回None
        """
        try:
            # 提取所有双引号内的内容
            quoted_strings = re.findall(r'"([^"]+)"', text)
            
            if quoted_strings:
                # 过滤出看起来像问题的字符串
                questions = []
                for q in quoted_strings:
                    q = q.strip()
                    # 验证：必须是有效的问题（长度足够且包含问号或以特定词开头）
                    if len(q) > 5 and (q.endswith('?') or '?' in q or 
                        any(kw in q for kw in ['如何', '为什么', '请描述', '请介绍', '请说明', '什么是', '谈谈你'])):
                        questions.append(q)
                
                if len(questions) >= 3:
                    return questions[:5]
            
            return None
        except Exception as e:
            print(f"提取引号内容失败: {str(e)}")
            return None
    
    def _parse_regex_questions(self, response):
        """
        使用正则表达式解析问题列表
        
        参数：
            response: LLM返回的响应文本
        
        返回：
            list: 问题列表
        """
        questions = []
        
        # 多种正则模式匹配问题
        patterns = [
            # 模式1: "问题1: 内容" 或 "问题1 内容"
            r'问题\s*(\d+)\s*[:：]\s*(.+?)(?=\n|$)',
            # 模式2: "1. 内容" 或 "1、内容"
            r'(?:^|\n)\s*(\d+)[.、]\s*(.+?)(?=\n|$)',
            # 模式3: "Q1: 内容"
            r'(?:^|\n)\s*Q\s*(\d+)\s*[:：]\s*(.+?)(?=\n|$)',
            # 模式4: 单独一行的问题（包含问号）
            r'(?:^|\n)\s*((?:问题|Q)?\s*\d*\s*[.、:：]?\s*.+\?)(?=\n|$)',
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, response, re.MULTILINE)
            if matches:
                for match in matches:
                    if isinstance(match, tuple):
                        # 提取问题内容（可能有两个分组）
                        question = match[1] if len(match) > 1 else match[0]
                    else:
                        question = match
                    
                    # 清理问题文本
                    question = question.strip()
                    # 移除可能的前缀序号
                    question = re.sub(r'^[\d\.\、\.\s\:\：]+', '', question)
                    
                    # 验证：必须是有效的问题（包含问号或以特定词开头）
                    if len(question) > 5 and (question.endswith('?') or '?' in question or 
                        any(kw in question for kw in ['如何', '为什么', '请描述', '请介绍', '请说明'])):
                        questions.append(question)
                
                # 如果找到足够的问题就停止
                if len(questions) >= 3:
                    break
        
        # 去重并限制数量
        unique_questions = []
        seen = set()
        for q in questions:
            # 使用问题前20个字符作为去重依据
            key = q[:20].lower()
            if key not in seen and len(q) > 5:
                seen.add(key)
                unique_questions.append(q)
        
        return unique_questions[:5]
    
    def _get_default_questions(self):
        """
        获取默认面试问题
        
        当 LLM 生成失败时使用
        """
        return [
            "请介绍一下你自己，以及你为什么对这个职位感兴趣？",
            "请详细描述一下你最成功的项目经历，你在其中扮演什么角色？",
            "你如何在团队中协作解决技术难题？请举例说明。",
            "如果遇到技术债务和业务需求冲突，你会如何平衡？",
            "你对自己未来3年的职业发展有什么规划？"
        ]
    
    def _display_current_question(self):
        """
        显示当前问题
        """
        # 检查面试是否已结束
        if self._interview_ended:
            print("\n⚠️ 面试已结束")
            return
        
        if self.current_question_idx < len(self.questions):
            q = self.questions[self.current_question_idx]
            remaining = len(self.questions) - self.current_question_idx
            print(f"\n📋 问题 {self.current_question_idx + 1}/{len(self.questions)} (还剩 {remaining} 题)：")
            print(f"   {q}")
        else:
            # 所有问题已显示完毕，标记结束并生成报告
            self._interview_ended = True
            self._generate_summary()
    
    def display_next_question(self):
        """
        显示下一个问题
        
        在 main.py 中调用，用于在提交回答后显示下一个问题
        """
        # 检查面试是否已结束
        if self._interview_ended:
            print("\n⚠️ 面试已结束")
            return
        
        if self.current_question_idx < len(self.questions):
            q = self.questions[self.current_question_idx]
            remaining = len(self.questions) - self.current_question_idx
            print(f"\n📋 问题 {self.current_question_idx + 1}/{len(self.questions)} (还剩 {remaining} 题)：")
            print(f"   {q}")
        else:
            # 所有问题已回答完毕，标记结束并生成报告
            self._interview_ended = True
            self._generate_summary()
    
    def submit_answer(self, answer):
        """
        提交候选人的回答
        
        参数：
            answer: 候选人的回答文本
        
        返回：
            dict: {
                'success': bool,             # 是否成功处理
                'need_follow_up': bool,    # 是否需要追问
                'follow_up_question': str,  # 追问问题（如果有）
                'is_interview_end': bool,  # 面试是否结束
                'evaluation': str          # 回答评价
            }
        
        示例：
            >>> agent.submit_answer("我之前在字节跳动负责...")
            {
                'success': True,
                'need_follow_up': False,
                'is_interview_end': False,
                'evaluation': '回答得很好！'
            }
        
        异常：
            RuntimeError: 如果面试尚未开始或已结束
            ValueError: 如果回答为空或无效
        """
        # ========== 严格的状态检查 ==========
        
        # 检查1：面试是否已开始
        if not hasattr(self, 'current_resume') or self.current_resume is None:
            raise RuntimeError("面试尚未开始，请先调用 start_interview()")
        
        # 检查2：面试是否已结束
        if self._interview_ended:
            raise RuntimeError("面试已结束，无法继续回答")
        
        # 检查3：问题列表是否有效
        if not self.questions or len(self.questions) == 0:
            raise RuntimeError("面试问题列表为空，面试无法进行")
        
        # 检查4：当前问题索引是否有效
        if self.current_question_idx < 0:
            raise RuntimeError("面试状态异常，问题索引为负数")
        
        # 检查5：是否已经回答完所有问题
        if self.current_question_idx >= len(self.questions):
            # 标记面试结束
            self._interview_ended = True
            print(f"\n{'='*60}")
            print(f"✅ 面试结束")
            print(f"   当前问题索引: {self.current_question_idx}, 总题数: {len(self.questions)}")
            print(f"{'='*60}")
            return {
                'success': False,
                'need_follow_up': False,
                'is_interview_end': True,
                'evaluation': '面试已结束'
            }
        
        # 检查6：回答是否有效
        if not answer or not isinstance(answer, str):
            raise ValueError("回答必须是非空字符串")
        
        # ========== 正常处理回答 ==========
        
        # 安全获取当前问题
        current_question = self.questions[self.current_question_idx]
        
        # 记录回答
        self.answers.append({
            'question_idx': self.current_question_idx,
            'question': current_question,
            'answer': answer.strip()
        })
        
        # 评价回答
        evaluation = self._evaluate_answer(current_question, answer)
        
        print(f"\n🤖 面试官评价：")
        print(f"   {evaluation}")
        
        # 移动到下一个问题
        self.current_question_idx += 1
        
        # 检查面试是否结束
        if self.current_question_idx >= len(self.questions):
            # 标记面试结束
            self._interview_ended = True
            return {
                'success': True,
                'need_follow_up': False,
                'is_interview_end': True,
                'evaluation': evaluation
            }
        
        # 不自动追问，直接进入下一个问题
        return {
            'success': True,
            'need_follow_up': False,
            'is_interview_end': False,
            'evaluation': evaluation
        }
    
    def _evaluate_answer(self, question, answer):
        """
        使用 LLM 评价候选人的回答（RAG 增强版）

        参数：
            question: 问题
            answer: 回答

        返回：
            str: 评价文本
        """
        # ====== RAG 增强：检索评分标准和参考答案 ======
        rag_context = ""
        try:
            rag = _get_rag()
            if rag:
                job_title = self.current_job.get('title', '') if self.current_job else ''
                rag_context = rag.retrieve_for_answer_evaluation(
                    question=question,
                    answer=answer,
                    job_title=job_title,
                    top_k=3,
                )
        except Exception as e:
            print(f"   [RAG] Evaluation context retrieval skipped: {e}")
            rag_context = ""
        # ====== RAG 增强结束 ======

        prompt = f"""作为面试官，请评价以下回答：

{rag_context}
【问题】
{question}

【回答】
{answer}

请从以下角度评价：
1. 回答的相关性和完整性
2. 技术深度和准确性
3. 表达清晰度

请用2-3句话给出简短评价，不要太长。
"""
        
        try:
            response = ask_glm(prompt)
            return response.strip()
        except Exception as e:
            return f"感谢您的回答。{len(answer) > 50 and '回答内容丰富' or '回答较为简洁'}。"
    
    def _should_follow_up(self, question, answer, evaluation):
        """
        判断是否需要追问（引入语义分析）
        
        参数：
            question: 问题
            answer: 回答
            evaluation: 评价
        
        返回：
            dict: {
                'need_follow_up': bool,     # 是否需要追问
                'confidence': float,        # 置信度 (0-1)
                'reason': str,              # 判断理由
                'suggestion': str           # 追问建议方向
            }
        """
        # 简单规则：回答太短时直接追问
        if len(answer.strip()) < 30:
            return {
                'need_follow_up': True,
                'confidence': 0.9,
                'reason': '回答过于简短',
                'suggestion': '请详细说明'
            }
        
        # 使用LLM进行语义分析，返回结构化判断结果
        prompt = f"""请分析以下问答，判断是否需要追问，并返回JSON格式结果：

【问题】
{question}

【回答】
{answer}

【初步评价】
{evaluation}

请判断：
1. 是否需要追问？（如果回答不够深入、不够具体、有歧义或需要更多细节）
2. 置信度是多少？（0-1，越高越确定）
3. 判断理由是什么？
4. 如果需要追问，建议追问方向是什么？

请返回JSON格式：
{{
    "need_follow_up": true或false,
    "confidence": 0.85,
    "reason": "判断理由",
    "suggestion": "追问建议方向（不需要则为空字符串）"
}}

注意：只输出JSON，不要有其他文字。
"""
        
        try:
            response = ask_glm(prompt)
            # 解析JSON响应
            result = self._parse_follow_up_response(response)
            if result is not None:
                return result
            
            # 如果JSON解析失败，使用简单规则
            return {
                'need_follow_up': len(answer.strip()) < 50,
                'confidence': 0.6,
                'reason': '解析失败，使用默认规则',
                'suggestion': ''
            }
        except Exception as e:
            print(f"警告：追问判断失败: {str(e)}")
            return {
                'need_follow_up': len(answer.strip()) < 50,
                'confidence': 0.5,
                'reason': 'LLM调用失败',
                'suggestion': ''
            }
    
    def _parse_follow_up_response(self, response):
        """
        解析追问判断的JSON响应
        
        参数：
            response: LLM返回的字符串
        
        返回：
            dict: 解析后的结果，如果失败返回None
        """
        try:
            # 尝试提取JSON部分
            if '```json' in response:
                json_str = response.split('```json')[1].split('```')[0].strip()
            elif '```' in response:
                json_str = response.split('```')[1].split('```')[0].strip()
            else:
                json_str = response.strip()
            
            data = json.loads(json_str)
            
            # 验证必要字段
            if isinstance(data, dict) and 'need_follow_up' in data:
                return {
                    'need_follow_up': bool(data['need_follow_up']),
                    'confidence': float(data.get('confidence', 0.5)),
                    'reason': str(data.get('reason', '')),
                    'suggestion': str(data.get('suggestion', ''))
                }
            
            return None
        except Exception as e:
            print(f"警告：解析追问响应失败: {str(e)}")
            return None
    
    def _generate_follow_up(self, question, answer):
        """
        生成追问问题
        
        参数：
            question: 原问题
            answer: 原回答
        
        返回：
            str: 追问问题
        """
        prompt = f"""基于以下问题和回答，请生成一个追问：

【问题】
{question}

【回答】
{answer}

请生成一个具体的追问，深入了解候选人的经验和能力。请直接输出追问问题，不要其他内容。
"""
        
        try:
            response = ask_glm(prompt)
            return response.strip()
        except:
            return "能否详细说明一下具体的技术细节？"
    
    def _generate_summary(self):
        """
        生成面试总结和评分报告
        
        返回：
            dict: {
                'summary_text': str,    # 完整的总结文本
                'overall_score': float, # 综合评分
                'strengths': str,       # 优势分析
                'weaknesses': str,      # 需要改进的地方
                'suggestions': str      # 录用建议
            }
        """
        print("\n" + "="*70)
        print("📊 面试评分与反馈报告")
        print("="*70)
        
        # 构建评分 prompt
        answers_text = "\n".join([
            f"问题{i+1}: {a['question']}\n回答: {a['answer']}"
            for i, a in enumerate(self.answers)
        ])
        
        total_questions = len(self.questions)
        answered_questions = len(self.answers)
        
        prompt = f"""请为这场面试生成一份详细的评分与反馈报告：

【候选人】{self.current_resume['name']}
【目标职位】{self.current_job['title']} @ {self.current_job['company']}
【完成情况】回答了 {answered_questions}/{total_questions} 个问题

【问答记录】
{answers_text}

请生成详细的面试报告，包括：

## 一、评分指标（请为每项打分 0-10，保留1位小数）
1. 技术能力评分：基于回答的技术深度和准确性
2. 表达沟通评分：基于回答的清晰度和逻辑性  
3. 经验匹配评分：基于经历与职位的相关度
4. 综合素质评分：综合整体表现

## 二、各项评分的详细说明
请对上述每项评分给出2-3句话的具体说明

## 三、优势分析
列出候选人的主要优势（3-5条）

## 四、需要改进的地方
列出需要改进的地方（3-5条）

## 五、综合评价与建议
给出是否推荐录用的建议，并说明理由

## 六、录用建议
- 推荐度：强烈推荐/推荐/一般/不推荐
- 建议薪资范围（如适用）
- 入职建议（如适用）

请用专业的语言回答，评分要客观公正。
"""
        
        try:
            summary = ask_glm(prompt)
            print(summary)
            
            # 提取综合评分（简单实现：从summary中提取第一个数字评分）
            import re
            # 查找评分数字（格式如 "8.5" 或 "8"）
            score_match = re.search(r'(\d+\.?\d*)', summary)
            overall_score = float(score_match.group(1)) if score_match else 0.0
            
            # 返回结构化数据
            return {
                'summary_text': summary,
                'overall_score': overall_score,
                'strengths': summary,  # 临时使用完整summary
                'weaknesses': summary,  # 临时使用完整summary
                'suggestions': summary  # 临时使用完整summary
            }
            
        except Exception as e:
            print(f"生成面试报告时出现错误: {e}")
            # 即使 LLM 调用失败，也生成基本信息
            fallback_summary = f"## 基础信息\n" \
                           f"- 候选人：{self.current_resume['name']}\n" \
                           f"- 目标职位：{self.current_job['title']}\n" \
                           f"- 完成度：{answered_questions}/{total_questions} 个问题\n" \
                           f"\n感谢您参加本次模拟面试！"
            print(fallback_summary)
            
            # 返回基础信息作为降级方案
            return {
                'summary_text': fallback_summary,
                'overall_score': 0.0,
                'strengths': '未生成',
                'weaknesses': '未生成',
                'suggestions': '未生成'
            }
        
        print("\n" + "="*70)
        print("✅ 面试结束")
        print("="*70)
    
    def end_interview(self):
        """
        提前结束面试
        
        当用户选择结束时调用此方法，生成面试报告
        """
        # 严格的状态检查
        if self._interview_ended:
            print("\n⚠️  面试已经结束，无法重复结束")
            return
        
        if not hasattr(self, 'current_resume') or self.current_resume is None:
            print("\n⚠️  面试尚未开始")
            return
        
        # 标记面试结束
        self._interview_ended = True
        
        answered = len(self.answers)
        remaining = len(self.questions) - answered
        
        print(f"\n⚠️  面试被提前结束")
        print(f"   已回答：{answered} 个问题")
        print(f"   剩余：{remaining} 个问题")
        
        # 生成面试报告
        self._generate_summary()
    
    def get_remaining_questions_count(self):
        """
        获取剩余问题数量

        返回：
            int: 剩余问题数量
        """
        return max(0, len(self.questions) - self.current_question_idx)


# =============================================================================
# 第三部分：面试报告生成工具
# =============================================================================

def generate_interview_report(session_data: dict, output_dir: str = "reports") -> str:
    """
    生成面试报告（Word格式）

    参数：
        session_data: 面试会话数据，包含以下字段：
            - candidate_name: 候选人姓名
            - job_title: 目标职位
            - company: 公司名称
            - questions: 问题列表
            - answers: 回答列表（包含 question、answer、evaluation、score）
            - summary: 面试总结
            - overall_score: 综合评分
            - interview_date: 面试日期（可选）
        output_dir: 输出目录，默认为 "reports"

    返回：
        str: 生成的报告文件路径

    示例：
        >>> session_data = {
        ...     'candidate_name': '张三',
        ...     'job_title': 'Python后端开发工程师',
        ...     'company': '创新科技有限公司',
        ...     'questions': ['问题1', '问题2', '问题3'],
        ...     'answers': [
        ...         {'question': '问题1', 'answer': '回答1', 'evaluation': '评价1', 'score': 8},
        ...         {'question': '问题2', 'answer': '回答2', 'evaluation': '评价2', 'score': 7},
        ...         {'question': '问题3', 'answer': '回答3', 'evaluation': '评价3', 'score': 9},
        ...     ],
        ...     'summary': '综合评价...',
        ...     'overall_score': 8
        ... }
        >>> generate_interview_report(session_data)
        'reports/面试报告_张三_20240115_1030.docx'
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.section import WD_ORIENT
        from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    except ImportError:
        raise ImportError("请安装 python-docx: pip install python-docx")

    # 确保输出目录存在
    os.makedirs(output_dir, exist_ok=True)

    # 获取面试日期
    interview_date = session_data.get('interview_date', '未知日期')
    
    # 生成文件名
    candidate_name = session_data.get('candidate_name', '未知候选人')
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    filename = f"面试报告_{candidate_name}_{timestamp}.docx"
    file_path = os.path.join(output_dir, filename)

    # 创建文档
    doc = Document()

    # 设置页面大小和边距
    section = doc.sections[0]
    section.top_margin = Pt(25.4)
    section.bottom_margin = Pt(25.4)
    section.left_margin = Pt(25.4)
    section.right_margin = Pt(25.4)

    # ================ 标题部分 ================
    title = doc.add_heading('AI 模拟面试报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(30, 58, 138)  # 深蓝色

    # 副标题
    subtitle = doc.add_paragraph(f"面试日期：{interview_date}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(11)
    subtitle.runs[0].font.color.rgb = RGBColor(107, 114, 128)

    doc.add_paragraph()  # 空行

    # ================ 基本信息表格 ================
    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = 'Table Grid'
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 设置表格样式
    for row in info_table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.width = Pt(180)

    # 填充表格内容
    info_table.cell(0, 0).text = '候选人姓名'
    info_table.cell(0, 1).text = candidate_name
    
    info_table.cell(1, 0).text = '目标职位'
    info_table.cell(1, 1).text = session_data.get('job_title', '未知职位')
    
    info_table.cell(2, 0).text = '应聘公司'
    info_table.cell(2, 1).text = session_data.get('company', '未知公司')

    # 设置表头样式
    for i in range(3):
        info_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
        info_table.cell(i, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(79, 70, 229)

    doc.add_paragraph()  # 空行

    # ================ 综合评分 ================
    overall_score = session_data.get('overall_score', 0)
    
    score_section = doc.add_heading('综合评分', level=1)
    score_section.runs[0].font.color.rgb = RGBColor(30, 58, 138)
    
    score_paragraph = doc.add_paragraph()
    score_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 大字号显示分数
    score_run = score_paragraph.add_run(f"{overall_score}/10")
    score_run.font.size = Pt(48)
    score_run.font.bold = True
    score_run.font.color.rgb = RGBColor(79, 70, 229)  # 靛蓝色

    # 添加评分等级
    if overall_score >= 9:
        level = '优秀'
        level_color = RGBColor(34, 197, 94)  # 绿色
    elif overall_score >= 7:
        level = '良好'
        level_color = RGBColor(34, 197, 94)  # 绿色
    elif overall_score >= 6:
        level = '合格'
        level_color = RGBColor(234, 179, 8)  # 黄色
    else:
        level = '需改进'
        level_color = RGBColor(239, 68, 68)  # 红色

    level_paragraph = doc.add_paragraph(f"评级：{level}")
    level_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    level_paragraph.runs[0].font.size = Pt(14)
    level_paragraph.runs[0].font.color.rgb = level_color

    doc.add_paragraph()  # 空行

    # ================ 问答记录 ================
    qa_section = doc.add_heading('问答记录', level=1)
    qa_section.runs[0].font.color.rgb = RGBColor(30, 58, 138)

    answers = session_data.get('answers', [])
    
    for idx, answer_data in enumerate(answers, 1):
        # 问题编号
        q_title = doc.add_heading(f"问题 {idx}", level=2)
        q_title.runs[0].font.color.rgb = RGBColor(59, 130, 246)  # 蓝色

        # 问题内容
        q_content = doc.add_paragraph(f"【问题】{answer_data.get('question', '')}")
        q_content.runs[0].font.size = Pt(11)

        # 回答内容
        a_content = doc.add_paragraph(f"【回答】{answer_data.get('answer', '')}")
        a_content.runs[0].font.size = Pt(11)
        a_content.runs[0].font.color.rgb = RGBColor(107, 114, 128)

        # 评价内容
        eval_content = doc.add_paragraph(f"【评价】{answer_data.get('evaluation', '')}")
        eval_content.runs[0].font.size = Pt(11)
        eval_content.runs[0].font.color.rgb = RGBColor(139, 92, 246)  # 紫色

        # 得分
        score = answer_data.get('score', 0)
        score_content = doc.add_paragraph(f"【得分】{score}/10")
        score_content.runs[0].font.size = Pt(11)
        score_content.runs[0].font.bold = True
        if score >= 8:
            score_content.runs[0].font.color.rgb = RGBColor(34, 197, 94)
        elif score >= 6:
            score_content.runs[0].font.color.rgb = RGBColor(234, 179, 8)
        else:
            score_content.runs[0].font.color.rgb = RGBColor(239, 68, 68)

        doc.add_paragraph()  # 空行

    # ================ 面试总结 ================
    summary_section = doc.add_heading('面试总结', level=1)
    summary_section.runs[0].font.color.rgb = RGBColor(30, 58, 138)

    summary_text = session_data.get('summary', '')
    if summary_text:
        # 尝试解析结构化总结
        try:
            if isinstance(summary_text, dict):
                # 结构化数据
                summary_dict = summary_text
                
                # 评分指标
                if 'scores' in summary_dict:
                    scores_heading = doc.add_heading('评分指标', level=2)
                    scores_table = doc.add_table(rows=len(summary_dict['scores']) + 1, cols=2)
                    scores_table.style = 'Table Grid'
                    scores_table.cell(0, 0).text = '评分项'
                    scores_table.cell(0, 1).text = '得分'
                    scores_table.cell(0, 0).paragraphs[0].runs[0].font.bold = True
                    scores_table.cell(0, 1).paragraphs[0].runs[0].font.bold = True
                    
                    for i, (key, value) in enumerate(summary_dict['scores'].items(), 1):
                        scores_table.cell(i, 0).text = key
                        scores_table.cell(i, 1).text = str(value)
                
                # 优势分析
                if 'strengths' in summary_dict:
                    strengths_heading = doc.add_heading('优势分析', level=2)
                    for strength in summary_dict['strengths']:
                        doc.add_paragraph(f"• {strength}")
                
                # 需要改进
                if 'improvements' in summary_dict:
                    improve_heading = doc.add_heading('需要改进', level=2)
                    for item in summary_dict['improvements']:
                        doc.add_paragraph(f"• {item}")
                
                # 综合评价
                if 'evaluation' in summary_dict:
                    eval_heading = doc.add_heading('综合评价', level=2)
                    doc.add_paragraph(summary_dict['evaluation'])
                
                # 录用建议
                if 'recommendation' in summary_dict:
                    rec_heading = doc.add_heading('录用建议', level=2)
                    doc.add_paragraph(f"推荐度：{summary_dict['recommendation']}")
            else:
                # 纯文本总结
                doc.add_paragraph(summary_text)
        except Exception as e:
            # 解析失败，直接显示文本
            doc.add_paragraph(summary_text)
    else:
        doc.add_paragraph("暂无详细总结")

    doc.add_paragraph()  # 空行

    # ================ 页脚 ================
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = "AI 模拟面试官系统生成"
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_paragraph.runs[0].font.size = Pt(10)
    footer_paragraph.runs[0].font.color.rgb = RGBColor(156, 163, 175)

    # 保存文档
    doc.save(file_path)

    print(f"✅ 面试报告已生成：{file_path}")
    return file_path


# =============================================================================
# 第四部分：LangChain 标准工具
# =============================================================================

# 尝试导入 LangChain tools（如果未安装则设为 None）
try:
    from langchain_core.tools import tool
    LANGCHAIN_AVAILABLE = True
except ImportError:
    # 如果没有安装 langchain_core，将 tool 定义为一个装饰器的替代版本
    LANGCHAIN_AVAILABLE = False
    def tool(func):
        """如果 LangChain 未安装，返回原始函数"""
        return func
    print("⚠️  提示：未安装 langchain-core，将使用基础函数模式。如需完整 LangChain 支持，请运行：pip install langchain-core")

@tool
def parse_resume_tool(file_path: str) -> Dict[str, Any]:
    """
    解析简历文件，提取结构化信息。
    
    这是 LangChain Agent 的简历解析工具，用于从 PDF、Word 或文本文件中
    提取候选人的关键信息，包括姓名、技能、教育背景、工作经验等。
    
    参数：
        file_path: 简历文件的完整路径（支持 .pdf, .docx, .doc, .txt 格式）
    
    返回：
        Dict[str, Any]: 包含以下字段的结构化简历数据：
            - name: 候选人姓名
            - skills: 技术技能列表
            - education: 教育背景
            - experience: 工作经验
            - text/raw_text: 原始文本内容
            - 其他从简历中提取的字段
    
    示例：
        >>> result = parse_resume_tool("/path/to/resume.pdf")
        >>> print(result['name'])
        '张三'
        >>> print(result['skills'])
        ['Python', 'SQL', '机器学习']
    
    注意事项：
        - 文件必须存在且格式受支持
        - 需要配置 ZHIPUAI_API_KEY 环境变量
    """
    try:
        parser = ResumeParser()
        result = parser.parse_file(file_path)
        
        # 确保返回的是字典格式
        if not isinstance(result, dict):
            result = {'raw_data': result}
        
        # 添加解析成功的元数据
        result['_parsed'] = True
        result['_file_path'] = file_path
        result['_parsed_at'] = datetime.now().isoformat()
        
        return result
        
    except FileNotFoundError as e:
        return {
            'error': f"文件未找到: {file_path}",
            '_parsed': False
        }
    except ValueError as e:
        return {
            'error': f"不支持的文件格式或解析失败: {str(e)}",
            '_parsed': False
        }
    except Exception as e:
        return {
            'error': f"简历解析失败: {str(e)}",
            '_parsed': False
        }


@tool
def evaluate_answer_tool(question: str, answer: str) -> str:
    """
    评价候选人对面试问题的回答。
    """
    prompt = f"""
你是一个严格的技术面试官。请评价候选人的回答。

面试问题：{question}
候选人的回答：{answer}

【评分标准】
- 1-3 分：回答完全错误、答非所问、或说"不知道/不清楚/不太会"
- 4-6 分：回答部分正确，但缺少关键细节、没有具体案例
- 7-8 分：回答基本正确，有核心要点，但深度不够
- 9-10 分：回答完整、有具体案例、有技术细节

输出格式（必须是 JSON）：
{{"score": 整数 1-10, "comment": "一句话评价"}}
"""
    try:
        response = ask_glm(prompt)
        # 解析 JSON
        import json
        import re
        json_match = re.search(r'\{.*\}', response, re.DOTALL)
        if json_match:
            result = json.loads(json_match.group())
            return json.dumps(result, ensure_ascii=False)
        return json.dumps({{"score": 5, "comment": "评价解析失败"}})
    except Exception as e:
        return json.dumps({{"score": 5, "comment": f"评价出错：{str(e)}"}})

@tool
def search_local_jobs(resume_data: dict) -> str:
    """
    搜索本地岗位数据库，根据简历数据中的技能匹配合适的岗位。
    
    这是 LangChain Agent 的岗位搜索工具，从本地 jobs.json 文件中读取岗位信息，
    根据简历数据中的技能列表进行匹配，返回所有匹配的岗位详情。
    
    参数：
        resume_data: 简历数据字典，必须包含 'skills' 字段
                    例如：{"name": "张三", "skills": ["Python", "FastAPI", "MySQL"], ...}
    
    返回：
        str: 匹配岗位的JSON格式字符串，包含岗位的完整信息。
             如果没有匹配的岗位，返回空数组的JSON字符串。
    
    匹配逻辑：
        - 从 resume_data 中提取 skills 字段
        - 遍历 jobs.json 中的所有岗位
        - 检查岗位的 requirements 字段是否包含用户的任一技能
        - 如果包含，则该岗位视为匹配
        - 返回所有匹配岗位的详细信息
    
    示例：
        >>> search_local_jobs({"name": "张三", "skills": ["Python", "Django"]})
        '[{"id": "1", "title": "Python后端开发工程师", ...}]'
    """
    import json
    import os
    
    # 从 resume_data 中提取技能列表
    skills = resume_data.get('skills', [])
    
    if not skills:
        return json.dumps({'error': '简历数据中没有技能信息，请先解析简历'})
    
    # 获取项目根目录下的 jobs.json 文件路径
    current_dir = os.path.dirname(os.path.abspath(__file__))
    jobs_file_path = os.path.join(current_dir, "jobs.json")
    
    try:
        # 读取 jobs.json 文件
        with open(jobs_file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        jobs = data.get('jobs', [])
        
        if not jobs:
            return json.dumps([])
        
        # 将技能转换为小写，便于比较
        skills_lower = [skill.lower() for skill in skills]
        
        # 匹配岗位
        matched_jobs = []
        for job in jobs:
            requirements = job.get('requirements', [])
            # 将岗位要求转换为字符串并转为小写
            requirements_str = ' '.join(requirements).lower()
            
            # 检查是否有任一技能匹配
            matched = False
            matched_skills = []
            for skill in skills_lower:
                if skill in requirements_str:
                    matched = True
                    matched_skills.append(skill)
            
            if matched:
                job['matched_skills'] = matched_skills
                matched_jobs.append(job)
        
        # 返回 JSON 格式的字符串
        return json.dumps(matched_jobs, ensure_ascii=False, indent=2)
        
    except FileNotFoundError:
        return json.dumps([{'error': f"文件未找到: {jobs_file_path}"}])
    except json.JSONDecodeError:
        return json.dumps([{'error': "jobs.json 文件格式错误"}])
    except Exception as e:
        return json.dumps([{'error': f"读取岗位信息失败: {str(e)}"}])

@tool
def get_current_time() -> str:
    """
    获取当前时间。
    当用户问“现在几点”或“当前时间”时调用此工具。
    """
    from datetime import datetime
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")

@tool
def get_weather(city: str) -> str:
    """
    获取指定城市的实时天气信息。
    当用户问“某地天气怎么样”、“某地热吗”、“某地冷不冷”时调用此工具。
    """
    import requests
    import json

    # 调用 高德 获取天气
    api_key = "614fe1567af465b360e014849a099638"
    base_url = f"https://restapi.amap.com/v3/weather/weatherInfo?city={city}&key={api_key}&extensions=base"
    try:
        response = requests.get(base_url, timeout=5)
        data = response.json()
        
        if data.get("status") == "1":
            live = data["lives"][0]
            weather = live.get("weather", "未知")
            temp = live.get("temperature", "未知")
            return f"{city}的天气：{weather}，{temp}度"
        else:
            return f"查询{city}天气失败：{data.get('info', '未知错误')}"
    except Exception as e:
        return f"查询天气失败：{str(e)}"

def format_job_results(matched_jobs_json: str) -> str:
    """
    将岗位匹配结果格式化为易读的文本格式
    
    参数：
        matched_jobs_json: search_local_jobs 返回的 JSON 字符串
    
    返回：
        格式化后的文本字符串
    """
    import json
    
    try:
        matched_jobs = json.loads(matched_jobs_json)
        
        if not matched_jobs:
            return "😔 抱歉，没有找到匹配的岗位。"
        
        if isinstance(matched_jobs, dict) and 'error' in matched_jobs:
            return f"❌ 错误: {matched_jobs['error']}"
        
        # 格式化输出
        result = []
        result.append(f"🎯 为您匹配到 {len(matched_jobs)} 个岗位\n")
        
        for idx, job in enumerate(matched_jobs, 1):
            title = job.get('title', '未知岗位')
            company = job.get('company', '未知公司')
            location = job.get('location', '未知地点')
            salary = job.get('salary', '面议')
            responsibilities = job.get('responsibilities', [])
            requirements = job.get('requirements', [])
            matched_skills = job.get('matched_skills', [])
            
            result.append(f"\n{'='*60}")
            result.append(f"📌 岗位 {idx}")
            result.append(f"🏢 {company} | 💼 {title} | 📍 {location} | 💰 {salary}")
            result.append(f"\n✅ 匹配技能: {', '.join(matched_skills)}")
            result.append(f"\n📋 职责描述:")
            for resp in responsibilities[:3]:  # 只显示前3条
                result.append(f"   • {resp}")
            result.append(f"\n📝 岗位要求:")
            for req in requirements[:4]:  # 只显示前4条
                result.append(f"   • {req}")
        
        result.append(f"\n{'='*60}")
        result.append("\n💡 以上是根据您的技能匹配的岗位，您可以选择感兴趣的进行深入了解！")
        
        return '\n'.join(result)
        
    except Exception as e:
        return f"❌ 格式化结果时出错: {str(e)}"


# =============================================================================
# 第五部分：RAG 知识库检索工具（LangChain Tool）
# =============================================================================

@tool
def search_knowledge_base(query: str) -> str:
    """
    搜索知识库，可以查询任何信息，包括：面试题、参考答案、评分标准、追问策略、候选人个人信息（如姓名、生日、学校）等。

    当需要以下信息时调用此工具：
    - 生成更有针对性的面试问题
    - 评估候选人回答时，需要参考答案或评分标准
    - 判断是否需要追问时，需要追问策略参考
    - 需要了解某个技术栈的常见面试考点

    参数：
        query: 搜索查询文本，可以是技术关键词或问题描述。例如：
               - "Python 后端 面试问题"
               - "FastAPI 依赖注入 评分标准"
               - "行为面试 STAR 法则"
               - "如何评价技术问题回答"

    返回：
        str: 格式化的相关知识内容，包含参考问题和答案
    """
    try:
        from knowledge_base import get_rag_chain
        rag = get_rag_chain()

        # 使用通用检索：从所有知识库中检索
        docs_with_scores = rag.kb.retrieve_with_scores(
            query=query,
            top_k=5,
        )

        if not docs_with_scores:
            return "[知识库] 未找到相关知识。请基于你的专业知识继续回答。"

        # 组装结果
        parts = ["[知识库检索结果]\n"]
        for i, (doc, score) in enumerate(docs_with_scores, 1):
            source = doc.metadata.get('source_file', '未知')
            section = doc.metadata.get('section', '')
            parts.append(f"--- 参考 {i}（来源：{source}）---")
            parts.append(doc.page_content.strip())
            parts.append("")

        parts.append("请基于以上知识库内容辅助你的判断和回答。")
        return "\n".join(parts)

    except Exception as e:
        return f"[知识库] 检索失败: {str(e)}。请基于你的专业知识继续。"

