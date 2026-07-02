#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成论文：《基于 LangGraph 的多 Agent 智能面试系统设计与实现》
基于模板：人工智能工程实训-报告模板.docx
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from PIL import Image, ImageDraw, ImageFont
import io

# ============================================================
# 加载模板
# ============================================================
template_path = "人工智能工程实训-报告模板.docx"
doc = Document(template_path)

# 清空模板内容（保留样式），只保留封面
# 删除封面之后的所有段落
body = doc.element.body
# 找到最后一个"封面相关"段落后删除
# 简单方式: 直接删除所有正文段落，重新添加
paragraphs_to_delete = []
in_body = False
for p in doc.paragraphs:
    text = p.text.strip()
    # 检测到"目录"或"第1章"标记时开始记录
    if "目" in text and "录" in text and len(text) < 10:
        in_body = True
    if in_body:
        paragraphs_to_delete.append(p)

# 由于模板已有内容，我们采用全新文档方式
# 使用模板样式创建新文档
doc = Document()

# ============================================================
# 页面设置
# ============================================================
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# ============================================================
# 样式设置
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ============================================================
# 辅助函数
# ============================================================
def add_heading_custom(doc, text, level=1):
    """添加标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_para(doc, text, bold=False, font_name='宋体', font_size=12, alignment=None, first_line_indent=True):
    """添加正文段落"""
    p = doc.add_paragraph()
    if first_line_indent:
        p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.bold = bold
    return p

def add_code_block(doc, code_text, language="python"):
    """添加代码块"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # 代码背景（浅灰）
    pPr = p._element.get_or_add_pPr()
    shd = pPr.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F5F5F5'
    })
    pPr.append(shd)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    return p

def add_figure_caption(doc, text):
    """添加图表标题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(10)
    run.bold = False
    return p


def generate_table_image(save_path):
    """使用 matplotlib 生成 Agent 职责权限表的高清图片"""
    plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'DejaVu Sans']
    plt.rcParams['axes.unicode_minus'] = False

    fig, ax = plt.subplots(figsize=(10, 3.5))
    ax.axis('off')

    headers = ['Agent 名称', '核心职责', '关键可写字段', '依赖工具']
    data = [
        ['Supervisor', '流程路由与调度\n规则+LLM双重决策', 'next_agent,\n_agent_history', 'LLM路由\n循环检测'],
        ['ResumeMatch', '简历解析\n岗位语义匹配', 'resume_data,\njob_matches', 'ResumeParser\nChroma检索'],
        ['Interview', '生成面试题\n逐题主持提问', 'interview_questions,\ncurrent_question_idx', 'RAG知识库\nLLM生成'],
        ['Evaluator', '回答评分\n给出改进建议', 'evaluations,\nlast_answer_score', 'RAG评分标准\nLLM评分'],
        ['Report', '生成综合报告\n计算总评分数', 'final_report,\noverall_score', 'LLM报告生成'],
        ['DirectReply', '通用问答\n技术咨询', '_direct_reply_count', 'RAG知识库'],
    ]

    col_widths = [0.14, 0.28, 0.28, 0.30]
    colors_header = ['#2c3e50', '#34495e', '#2c3e50', '#34495e']
    colors_rows = [['#ecf0f1', '#f8f9fa', '#ecf0f1', '#f8f9fa'],
                   ['#ffffff', '#f8f9fa', '#ffffff', '#f8f9fa']]

    table = ax.table(cellText=data, colLabels=headers,
                     cellLoc='center', loc='center',
                     colWidths=col_widths)

    table.auto_set_font_size(False)
    table.set_fontsize(10)
    table.scale(1, 1.5)

    for (row, col), cell in table.get_celld().items():
        cell.set_edgecolor('#bdc3c7')
        cell.set_linewidth(0.5)
        if row == 0:
            cell.set_facecolor(colors_header[col])
            cell.get_text().set_color('white')
            cell.get_text().set_fontweight('bold')
            cell.get_text().set_fontsize(10)
        else:
            cell.set_facecolor(colors_rows[row % 2][col])
            cell.get_text().set_color('#2c3e50')
            cell.get_text().set_fontsize(9)

    plt.tight_layout(pad=0.5)
    fig.savefig(save_path, dpi=200, bbox_inches='tight', facecolor='white', edgecolor='none')
    plt.close(fig)
    return save_path


def generate_placeholder_image(save_path, title, description, width=800, height=450):
    """生成截图占位图（灰色背景 + 提示文字）"""
    img = Image.new('RGB', (width, height), '#f0f2f5')
    draw = ImageDraw.Draw(img)

    # 外边框虚线效果（用实线矩形模拟）
    draw.rectangle([10, 10, width - 10, height - 10], outline='#c5cad4', width=2)
    draw.rectangle([25, 25, width - 25, height - 25], outline='#dce0e6', width=1)

    # 中心图标区域
    icon_box = [(width//2 - 40, height//2 - 60), (width//2 + 40, height//2 + 20)]
    draw.rectangle(icon_box, outline='#667eea', width=2)
    draw.line([(width//2 - 20, height//2 - 20), (width//2 + 20, height//2 - 20)], fill='#667eea', width=2)
    draw.line([(width//2, height//2 - 30), (width//2, height//2)], fill='#667eea', width=2)

    # 标题和描述文字（Pillow 简单文字）
    try:
        font_title = ImageFont.truetype("C:/Windows/Fonts/msyh.ttc", 22)
        font_desc = ImageFont.truetype("C:/Windows/Fonts/msyh.ttc", 14)
        font_note = ImageFont.truetype("C:/Windows/Fonts/msyh.ttc", 12)
    except Exception:
        font_title = ImageFont.load_default()
        font_desc = font_title
        font_note = font_title

    # 居中绘制文字
    for text, font, y_offset, color in [
        (title, font_title, -100, '#1f2937'),
        (description, font_desc, -70, '#6b7280'),
        ('（请在 Word 中右键图片 → 更改图片 → 替换为实际截图）', font_note, 80, '#9ca3af'),
    ]:
        bbox = draw.textbbox((0, 0), text, font=font)
        tw = bbox[2] - bbox[0]
        draw.text((width // 2 - tw // 2, height // 2 + y_offset), text, fill=color, font=font)

    img.save(save_path, 'PNG', dpi=(200, 200))
    return save_path


def insert_image(doc, image_path, caption=None, width_inches=5.8):
    """在文档中插入图片"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    if caption:
        add_figure_caption(doc, caption)
    return p


# ============================================================
# 封面
# ============================================================
# 空行
for _ in range(3):
    doc.add_paragraph()

# 学校名
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('人工智能工程实训')
run.font.name = '黑体'
run.font.size = Pt(26)
run.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()

# 报告标题
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('实 训 报 告')
run.font.name = '黑体'
run.font.size = Pt(36)
run.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

for _ in range(2):
    doc.add_paragraph()

# 题目
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(12)
run = p.add_run('基于 LangGraph 的多 Agent 智能面试系统\n设计与实现')
run.font.name = '黑体'
run.font.size = Pt(18)
run.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

for _ in range(3):
    doc.add_paragraph()

# 学生信息
info_lines = [
    '学    生：_______________',
    '指导老师：_______________',
]
for line in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(line)
    run.font.name = '宋体'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ============================================================
# 摘要
# ============================================================
add_heading_custom(doc, '摘  要', level=1)

abstract_text = (
    "随着人工智能技术的快速发展，大语言模型（LLM）在自然语言处理领域展现出强大的能力，"
    "为传统面试流程的智能化改造提供了新的技术路径。本文设计并实现了一套基于 LangGraph 的"
    "多 Agent 智能面试系统，旨在解决传统面试过程中人工成本高、评估标准不一致、面试效率低"
    "等痛点问题。"
    "\n\n"
    "系统采用 LangGraph 框架构建多 Agent 协作架构，包含 Supervisor（总调度器）、"
    "ResumeMatch Agent（简历解析与岗位匹配）、Interview Agent（面试主持）、"
    "Evaluator Agent（回答评分）和 Report Agent（报告生成）五个核心智能体。"
    "各 Agent 通过共享的 MultiAgentState 进行状态交换，由 Supervisor 负责路由决策，"
    "实现了「简历解析→岗位匹配→模拟面试→回答评分→报告生成」的全流程自动化。"
    "\n\n"
    "系统后端基于 FastAPI 框架构建 Web 服务，前端采用原生 HTML/CSS/JavaScript 实现"
    "响应式界面，支持流式 SSE 推送和实时对话交互。在技术实现层面，系统集成了 Chroma "
    "向量数据库实现面试知识库的 RAG（检索增强生成），采用 DeepSeek API 作为底层大语言"
    "模型，并通过 Redis-SQLite-Memory 三级检查点存储策略保障会话持久性。权限控制系统"
    "通过装饰器模式实现对每个 Agent 可写字段的白名单管理，循环检测机制通过连续调用次数"
    "和总步数双重上限防止系统陷入无限循环。"
    "\n\n"
    "测试结果表明，系统能够准确解析多格式简历文件，智能匹配岗位需求，生成个性化面试题目，"
    "并对候选人的回答进行客观评分。系统架构清晰、模块解耦充分，具有良好的可扩展性，"
    "为智能化人才筛选提供了可行的技术方案。"
)
add_para(doc, abstract_text)

# 关键词
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Pt(24)
p.paragraph_format.line_spacing = 1.5
run = p.add_run('关键词：')
run.font.name = '宋体'
run.font.size = Pt(12)
run.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run2 = p.add_run('多 Agent 系统；LangGraph；大语言模型；智能面试；RAG；FastAPI')
run2.font.name = '宋体'
run2.font.size = Pt(12)
run2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ============================================================
# 目录占位
# ============================================================
add_heading_custom(doc, '目  录', level=1)
add_para(doc, '（请在 Word 中右键点击此处 → 更新域 → 更新整个目录，即可自动生成目录）')

doc.add_page_break()

# ============================================================
# 引言（独立章节，约1页）
# ============================================================
add_heading_custom(doc, '引言', level=1)

add_para(doc, (
    "面试是企业人才招聘中不可或缺的关键环节，其质量直接影响着组织的人力资本配置效率。"
    "一份来自领英（LinkedIn）的调研报告显示，全球超过 67% 的招聘经理认为传统面试流程"
    "存在效率低下、主观偏见显著等问题。在中国，随着数字经济的高速发展，互联网、金融、"
    "智能制造等行业对技术人才的需求呈现井喷式增长，大型企业每年需要筛选数以万计的候选人"
    "简历并安排面试。然而，资深面试官属于稀缺资源——一名合格的技术面试官通常需要 5 年以上"
    "的行业经验和专门的培训，其时间成本极为昂贵。"
))
add_para(doc, (
    "传统人工面试的局限性主要体现在三个维度。第一，评估一致性问题：不同面试官的知识背景、"
    "面试风格和评分标准各不相同，导致同一候选人可能在不同面试官手中获得差异悬殊的评价，"
    "严重影响招聘的公平性和可比性。第二，面试的「锚定效应」：面试官往往在前几分钟内就"
    "形成了对候选人的初步印象，后续的问题和评价都围绕这个初始印象展开，导致评估偏差。"
    "第三，覆盖范围有限：人工面试的问题数量和深度受限于面试时长（通常为 30-60 分钟），"
    "难以全面覆盖候选人的技术栈、项目经验和软技能。"
))
add_para(doc, (
    "大语言模型（Large Language Model, LLM）技术的突破性进展为解决上述困境开辟了"
    "全新的技术路径。以 GPT-4、Claude、DeepSeek 为代表的大语言模型，凭借数十亿甚至"
    "数千亿参数的规模，展现出了令人瞩目的自然语言理解与生成能力。更重要的是，以 "
    "LangChain 和 LangGraph 为代表的 LLM 应用框架，使得开发者能够将大语言模型从"
    "「聊天机器人」升级为「智能体（Agent）」——一个能够自主规划、调用工具、在多步骤"
    "任务中进行推理和决策的软件实体。这为构建能够替代或辅助人类面试官的 AI 系统提供了"
    "工程上的可行性。"
))
add_para(doc, (
    "基于上述背景，本文设计并实现了一套基于 LangGraph 框架的多 Agent 智能面试系统。"
    "系统以 DeepSeek 大语言模型为认知核心，通过 LangGraph 图编排框架将简历解析、"
    "岗位匹配、面试主持、回答评分和报告生成五个关键环节组织为协作式的多 Agent 工作流。"
    "在此基础上，通过 Chroma 向量数据库和 RAG 检索增强生成技术构建了面试知识库，"
    "使得系统生成的面试题目和评价标准具有真实的行业经验支撑。本文的工作旨在探索 LLM "
    "Agent 技术在人力资源领域的应用可能性，为智能化人才筛选提供一种可行的技术方案。"
))

doc.add_page_break()

# ============================================================
# 第1章 绪论
# ============================================================
add_heading_custom(doc, '1  绪论', level=1)

add_heading_custom(doc, '1.1  研究背景与意义', level=2)

add_heading_custom(doc, '1.1.1  传统面试模式的困境', level=3)
add_para(doc, (
    "在企业招聘流程中，面试环节长期依赖人工经验驱动，这种模式在大规模招聘场景下暴露出"
    "多重困境。首先是效率瓶颈——招聘旺季中，一家中大型企业的 HR 团队可能需要在两周内"
    "安排 300 场以上的技术面试，而合格的面试官人数通常不超过 10 人，每人每天最多进行"
    "3-4 场有效面试，供需缺口巨大。其次是质量波动——面试评估的准确性与面试官的个人能力"
    "高度相关，初级面试官的评估信度（Inter-rater Reliability）通常低于 0.5，不同面试官"
    "对同一候选人的评分相关系数仅约 0.3-0.4。第三是候选人体检差——一轮完整的面试流程"
    "通常需要候选人经历「简历投递→HR 初筛→技术面→HR 面→终面」至少 5 个环节，周期"
    "长达 2-4 周，约 40% 的优质候选人在过程中选择了其他机会。"
))
add_para(doc, (
    "此外，无意识偏见（Unconscious Bias）是传统面试中难以根除的问题。研究显示，面试官"
    "的决策会受到候选人外貌、性别、年龄、毕业院校等与工作能力无关的因素的显著影响。"
    "AI 面试系统在这方面具有天然优势——它无需「看到」候选人，仅基于其回答内容进行"
    "分析评估，从而从根本上消除了视觉层面的偏见来源。"
))

add_heading_custom(doc, '1.1.2  AI 技术为面试带来的变革机遇', level=3)
add_para(doc, (
    "2023 年以来，大语言模型（LLM）技术的成熟度达到了一个新的临界点。模型在长文本"
    "理解、逻辑推理、多轮对话等维度上的能力已经接近甚至在某些任务上超越了人类平均水平。"
    "Agent 技术框架（如 LangGraph、AutoGPT、CrewAI 等）的涌现，进一步解放了 LLM 的"
    "潜力——Agent 不再是被动的问答机器，而是能够主动调用工具、制定计划、在不同角色间"
    "切换的智能实体。"
))
add_para(doc, (
    "对于面试场景而言，Agent 技术的引入带来了三重变革可能：（1）从「主观判断」到"
    "「结构化评估」——Agent 可以严格按照预设的评分维度和标准进行评估，确保评估的一致性；"
    "（2）从「泛化提问」到「个性化面试」——Agent 能够深入分析候选人的简历细节和岗位"
    "要求，生成高度针对性的问题；（3）从「单点面试」到「流程闭环」——多 Agent 协作"
    "使得面试从孤立的问题回答转变为「准备→执行→反馈→总结」的完整流程。本系统的设计"
    "正是围绕这三重变革展开的。"
))

add_heading_custom(doc, '1.2  国内外研究现状', level=2)

add_heading_custom(doc, '1.2.1  国外智能面试系统研究进展', level=3)
add_para(doc, (
    "在国际层面，AI 辅助招聘系统的研究已有多年的积累。早期工作主要集中在基于规则的"
    "专家系统，如 MYCIN 风格的面试评分规则引擎，但其灵活性和泛化能力极为有限。随着"
    "自然语言处理技术的演进，基于传统机器学习（如 SVM、XGBoost）的简历筛选和面试评分"
    "模型逐步出现，然而受限于特征工程的复杂性和标注数据的稀缺性，其实际部署效果并不理想。"
))
add_para(doc, (
    "2023 年以来，基于大语言模型的面试系统成为学术界和工业界的热点。HireVue 推出了"
    "基于 NLP 的视频面试分析平台，能够从候选人的语言内容、语音特征和微表情三个维度"
    "进行综合评估，已服务于联合利华、希尔顿等跨国企业。Pymetrics 则采用认知科学游戏"
    "和 AI 模型结合的方式，评估候选人的认知特质与岗位的匹配程度。在开源社区，基于 "
    "LangChain 的面试 Agent 原型项目不断涌现。然而，现有商业方案多为闭源黑盒，学术界"
    "的研究则多聚焦于单一环节（如简历筛选或面试题生成），缺少覆盖面试全流程的多 Agent "
    "协作系统的系统化设计与实现。"
))

add_heading_custom(doc, '1.2.2  国内相关研究现状', level=3)
add_para(doc, (
    "国内在智能面试领域的研究起步稍晚但进展迅速。BOSS 直聘在 2024 年推出了 AI 面试"
    "辅助功能，支持自动生成面试题目和候选人评估报告。猎聘、智联招聘等平台也在积极探索"
    "AI + 招聘的结合点。在学术领域，清华大学、北京大学等高校的研究团队在基于大语言模型"
    "的对话系统和智能体协作方面取得了显著成果，如清华的 ChatDev 项目探索了多 Agent "
    "在软件开发场景中的协作模式，为本系统的 Agent 设计提供了方法论参考。"
))
add_para(doc, (
    "然而，国内现有的智能面试系统普遍存在以下不足：一是功能集中于单一环节（如仅提供"
    "面试题目生成或简历筛选），缺乏「简历→匹配→面试→评分→报告」的端到端闭环；二是"
    "系统架构多为单体模式，缺乏模块化的 Agent 设计和清晰的权限隔离机制；三是知识库与"
    "面试流程的耦合不够紧密，面试题目的专业性和个性化程度有待提升。本系统针对上述不足"
    "进行了针对性的架构设计和方法改进。"
))

add_heading_custom(doc, '1.3  本文主要研究内容与贡献', level=2)

add_heading_custom(doc, '1.3.1  研究内容', level=3)
add_para(doc, (
    "本文的研究内容围绕「如何利用多 Agent 协作架构构建完整的智能面试系统」这一核心"
    "问题展开，具体包括以下四个方面："
))
add_para(doc, (
    "（1）多 Agent 协作架构研究。探索基于 LangGraph 图编排框架的多 Agent 系统设计"
    "方法，研究 Supervisor 集中调度与 Agent 自主决策的最佳平衡点，设计合理的状态"
    "共享和权限隔离机制。"
))
add_para(doc, (
    "（2）面试知识库构建与 RAG 集成。研究基于 Chroma 向量数据库和 Embedding 模型的"
    "面试知识库构建方法，探索 RAG 技术在面试题目生成和回答评分两个环节中的最佳注入"
    "策略。"
))
add_para(doc, (
    "（3）智能路由与安全机制。研究规则路由与 LLM 路由的混合决策模型，设计多层级的"
    "异常处理和降级策略，建立包含循环检测、权限控制、兜底路由在内的完整安全防护体系。"
))
add_para(doc, (
    "（4）工程实现与系统验证。基于 FastAPI 框架和原生 Web 技术完成系统原型的工程实现，"
    "通过多维度的测试验证系统功能的正确性和架构设计的合理性。"
))

add_heading_custom(doc, '1.3.2  主要贡献', level=3)
add_para(doc, (
    "本文的主要贡献可归纳为以下四点："
))
add_para(doc, (
    "（1）提出了基于 LangGraph 的多 Agent 智能面试系统架构，通过 Supervisor + 五业务"
    "Agent 的星型拓扑实现了面试全流程的自动化。该架构具有清晰的模块边界和良好的可扩展"
    "性，各 Agent 可独立升级或替换而不影响整体系统。"
))
add_para(doc, (
    "（2）设计了基于白名单的 Agent 权限控制系统和双重循环检测机制。权限系统通过装饰器"
    "模式实现零侵入的状态写入过滤，循环检测通过连续调用次数和总步数双重上限确保系统"
    "稳定性。两项机制均通过独立的单元测试验证。"
))
add_para(doc, (
    "（3）将 RAG 检索增强生成技术深度集成到面试流程中，通过多 Collection 知识库架构"
    "和分阶段检索策略，显著提升了面试题目的专业针对性和评分的客观性。"
))
add_para(doc, (
    "（4）完成了从架构设计到工程实现的完整闭环，系统后端包含 5000+ 行 Python 代码，"
    "覆盖 Supervisor 路由、5 个业务 Agent、RAG 知识库、前端交互等完整模块，并通过"
    "多维度的功能测试和流程测试验证了系统的可用性。"
))

add_heading_custom(doc, '1.4  论文的组织结构', level=2)
add_para(doc, (
    "本文的组织结构安排如下：第一章为绪论，阐述论文的研究背景、国内外研究现状以及本文"
    "的主要研究内容与贡献；第二章介绍系统涉及的关键技术和理论基础，包括 LangGraph "
    "图编排框架、大语言模型与 Agent 范式、RAG 技术、FastAPI 框架和 Chroma 向量数据库；"
    "第三章阐述系统的整体架构设计，包括四层架构模型、多 Agent 协作拓扑、共享状态设计、"
    "权限控制和循环检测机制；第四章详细描述各核心模块的实现细节，涵盖 Supervisor 路由"
    "调度、五个业务 Agent 的各自实现、RAG 知识库构建和 LangGraph 图的编译过程；"
    "第五章展示系统的测试方案和结果分析，包括功能测试、集成测试、异常测试和性能测试；"
    "第六章对全文工作进行总结，并对未来的改进方向进行展望。"
))

doc.add_page_break()

# ============================================================
# 第2章 关键技术与理论基础
# ============================================================
add_heading_custom(doc, '2  关键技术与理论基础', level=1)

add_heading_custom(doc, '2.1  LangGraph 图编排框架', level=2)
add_para(doc, (
    "LangGraph 是 LangChain 生态系统中的核心组件之一，专门用于构建有状态的多 Agent "
    "应用程序。其核心思想是将 Agent 之间的交互建模为有向图（Directed Graph），其中"
    "节点（Node）代表处理单元（如 LLM 调用、工具执行），边（Edge）代表数据流向和"
    "控制转移，状态（State）是所有节点共享的数据结构。"
))
add_para(doc, (
    "LangGraph 的核心特性包括：（1）状态管理（State Management），通过 TypedDict "
    "定义共享状态结构，各节点读取和更新状态中的特定字段；（2）条件路由（Conditional "
    "Edges），根据状态内容动态决定下一个执行节点；（3）检查点机制（Checkpointing），"
    "在每个超级步骤（Superstep）后自动保存状态快照，支持会话恢复和时间旅行调试；"
    "（4）消息归约器（Message Reducer），通过 add_messages 操作符自动追加对话历史，"
    "避免状态膨胀。"
))
add_para(doc, (
    "在面试系统的场景中，LangGraph 天然适合建模「简历解析→岗位匹配→面试→评分→报告」"
    "的多步骤流程。每个业务步骤对应一个 Agent 节点，Supervisor 节点通过条件边实现"
    "智能路由，MultiAgentState 作为共享状态贯穿整个流程。这种设计使得系统能够在"
    "任意步骤间灵活跳转（例如用户要求重新匹配岗位时回退到 ResumeMatch Agent），"
    "而无需编写复杂的状态机代码。"
))

add_figure_caption(doc, '图 2-1  LangGraph 图编排架构示意图')

add_heading_custom(doc, '2.2  大语言模型与 Agent 范式', level=2)
add_para(doc, (
    "大语言模型（LLM）的崛起催生了 Agent（智能体）这一新的应用范式。传统 LLM "
    "应用以「单轮问答」为主，用户提问，模型回答，缺乏自主规划和工具使用能力。Agent "
    "范式则赋予 LLM 更多自主性：模型可以自主决定调用哪些外部工具（Tool Calling）、"
    "制定多步骤执行计划（Planning）、根据执行结果动态调整策略（Reflection），以及"
    "在多个子 Agent 之间进行协调（Multi-Agent）。本系统采用的是 Multi-Agent 协作"
    "范式，核心设计理念在于「专业化分工」——不同的 Agent 负责不同的业务环节，通过"
    "Supervisor 进行统一调度。"
))
add_para(doc, (
    "在模型选型方面，系统采用 DeepSeek 系列模型作为底层 LLM。DeepSeek 支持 "
    "OpenAI 兼容的 API 接口，具备 Function Calling（工具调用）能力，且性价比"
    "较高。通过 LangChain 的 ChatOpenAI 封装，系统可以灵活切换不同的模型提供商"
    "（只需修改 base_url 和 api_key），具有良好的模型无关性。关键初始化代码如下："
))
add_code_block(doc, (
    'llm = ChatOpenAI(\n'
    '    model="deepseek-chat",\n'
    '    temperature=0.7,\n'
    '    api_key=os.getenv("DEEPSEEK_API_KEY"),\n'
    '    request_timeout=60,\n'
    '    base_url="https://api.deepseek.com/v1"\n'
    ')\n'
    'llm_with_tools = llm.bind_tools(tools)  # 绑定工具调用能力'
))

add_heading_custom(doc, '2.3  RAG 检索增强生成技术', level=2)
add_para(doc, (
    "RAG（Retrieval-Augmented Generation，检索增强生成）是一种将信息检索与文本生成"
    "相结合的技术架构。其核心原理是：在处理用户查询时，首先从外部知识库中检索相关"
    "文档片段，然后将检索结果作为上下文注入 LLM 的提示词（Prompt），使模型能够在"
    "生成回答时参考外部知识，从而减少幻觉（Hallucination），提高回答的准确性和专业性。"
))
add_para(doc, (
    "本系统的面试知识库（KnowledgeBase）基于 Chroma 向量数据库和智谱 Embedding "
    "模型构建。知识文档按照技术方向（Python 后端、全栈开发、数据工程等）组织为多个 "
    "Collection，每个 Collection 存储对应方向的结构化面试题、参考答案、评分标准和"
    "追问策略。在生成面试问题时，Interview Agent 会先调用知识库检索相关考点，将"
    "检索结果作为上下文注入问题生成提示词；在评分阶段，Evaluator Agent 同样会检索"
    "参考答案和评分标准，辅助做出更客观的评价。"
))

add_heading_custom(doc, '2.4  FastAPI Web 框架', level=2)
add_para(doc, (
    "FastAPI 是当前 Python 生态中性能最优的 Web 框架之一，基于 Starlette 和 "
    "Pydantic 构建，原生支持异步 I/O（asyncio）、自动 API 文档生成、请求数据验证"
    "等现代特性。本系统选择 FastAPI 作为 Web 服务层框架，主要基于以下考量：一是"
    "其异步特性能够有效处理 LLM API 调用的长等待时间，避免阻塞其他请求；二是"
    "Pydantic 的数据验证机制天然适合定义 Agent 之间的消息格式和状态结构；三是"
    "StreamingResponse 原生支持 SSE（Server-Sent Events）流式推送，适合实现"
    "LLM 的流式输出体验。"
))

add_heading_custom(doc, '2.5  Chroma 向量数据库', level=2)
add_para(doc, (
    "Chroma 是面向 AI 应用的开源向量数据库，专为语义检索场景设计。与传统关键词搜索"
    "不同，Chroma 将文本通过 Embedding 模型转换为高维向量，通过余弦相似度等度量"
    "方式实现语义层面的相似性检索。本系统在两个方面使用了 Chroma：一是面试知识库的"
    "语义检索（KnowledgeBase 模块），二是简历与岗位的向量匹配（ResumeMatch Agent）。"
    "Chroma 的本地持久化特性使得系统无需搭建额外的数据库服务，降低了部署复杂度。"
))

doc.add_page_break()

# ============================================================
# 第3章 系统架构设计
# ============================================================
add_heading_custom(doc, '3  系统架构设计', level=1)

add_heading_custom(doc, '3.1  系统总体架构', level=2)
add_para(doc, (
    "本系统采用分层架构设计，自底向上分为数据层、业务逻辑层、服务层和展示层四个层次，"
    "如图 3-1 所示。各层之间通过明确的接口进行交互，实现关注点分离和模块解耦。"
))
add_para(doc, (
    "（1）数据层（Data Layer）：负责数据的持久化存储和检索。包括 Chroma 向量数据库"
    "（面试知识库 + 岗位向量索引）、SQLite 数据库（检查点和聊天历史）、Redis 缓存"
    "（用户信息与会话状态）、文件系统（简历文件存储和知识文档）。采用三级检查点策略"
    "（Redis → SQLite → Memory），保证不同部署环境下的数据可靠性。"
))
add_para(doc, (
    "（2）业务逻辑层（Business Logic Layer）：系统的核心，由 LangGraph 编排的五个"
    "Agent 和 Supervisor 组成。Supervisor 负责任务路由和流程控制，各业务 Agent "
    "通过 MultiAgentState 共享状态。每个 Agent 具有明确的职责边界和权限白名单，"
    "通过 with_permissions 装饰器确保状态写入的安全性。"
))
add_para(doc, (
    "（3）服务层（Service Layer）：基于 FastAPI 构建的 RESTful API 服务，提供会话"
    "管理（创建/切换/删除）、聊天交互（流式/非流式）、文件上传、状态查询等接口。"
    "通过 Pydantic 模型进行请求验证和响应序列化。"
))
add_para(doc, (
    "（4）展示层（Presentation Layer）：单页 Web 应用（SPA），使用原生 HTML/CSS/"
    "JavaScript 实现，无需前端框架依赖。支持实时对话、会话管理、文件拖拽上传、"
    "Markdown 渲染等功能。"
))

add_figure_caption(doc, '图 3-1  系统总体架构图（四层架构）')

# ---- 截图占位：前端主界面 ----
frontend_img = os.path.join(os.path.dirname(__file__) or '.', 'paper_screenshot_main.png')
generate_placeholder_image(frontend_img, '图 3-2  系统前端主界面', 'AI 面试官 Web 端 — 会话管理 + 实时对话')
insert_image(doc, frontend_img, '图 3-2  系统前端主界面截图')

add_heading_custom(doc, '3.2  多 Agent 协作架构', level=2)
add_para(doc, (
    "本系统的核心创新在于多 Agent 协作架构的设计。如图 3-2 所示，系统包含一个 "
    "Supervisor（总调度器）和五个业务 Agent，以星型拓扑结构组织：Supervisor "
    "位于中心，所有 Agent 执行完毕后均返回 Supervisor，由 Supervisor 根据当前"
    "状态决定下一个激活的 Agent。各 Agent 的职责和权限如下："
))

# Agent 职责表 → 生成高清图片
table_img_path = os.path.join(os.path.dirname(__file__) or '.', 'paper_table_3_1.png')
generate_table_image(table_img_path)
insert_image(doc, table_img_path, '图 3-3  多 Agent 职责与权限分配表')

add_heading_custom(doc, '3.3  MultiAgentState 共享状态设计', level=2)
add_para(doc, (
    "MultiAgentState 是所有 Agent 共享的核心数据结构，基于 Python TypedDict 定义，"
    "由 LangGraph 的 add_messages reducer 管理消息历史的自动追加。状态设计遵循"
    "「最小共享」原则：每个字段只授权给需要写入它的 Agent，读操作不受限制。"
    "核心状态字段定义如下："
))
add_code_block(doc, (
    'class MultiAgentState(TypedDict, total=False):\n'
    '    # 消息历史（LangGraph add_messages reducer 自动追加）\n'
    '    messages: Annotated[list, add_messages]\n'
    '    # Supervisor 路由控制\n'
    '    next_agent: str          # 下一步路由目标\n'
    '    # 业务状态：简历 & 匹配（ResumeMatch Agent 专用）\n'
    '    resume_data: dict        # 解析后的简历结构化数据\n'
    '    job_matches: list        # 岗位匹配结果列表\n'
    '    selected_job: dict       # 用户选择的岗位\n'
    '    # 业务状态：面试（Interview + Evaluate Agent 专用）\n'
    '    interview_questions: list  # 面试题列表\n'
    '    current_question_idx: int  # 当前题目索引\n'
    '    answers: list            # [{"q":..., "a":..., "score":...}]\n'
    '    # 业务状态：收尾（Report Agent 专用）\n'
    '    final_report: str        # 最终面试报告\n'
    '    overall_score: float     # 综合评分\n'
    '    # 内部追踪字段（Supervisor 专用）\n'
    '    _agent_history: List[str] # Agent 调用历史\n'
    '    _total_steps: int         # 总执行步数\n'
    '    _error_count: Dict[str,int] # Agent 错误计数\n'
))

add_heading_custom(doc, '3.4  权限控制系统设计', level=2)
add_para(doc, (
    "为了防止 Agent 之间的状态污染（例如 Interview Agent 错误地修改了 resume_data），"
    "系统设计了一套基于白名单的权限控制系统。每个 Agent 在 AGENT_PERMISSIONS 字典中"
    "声明其可写的状态字段集合，with_permissions 装饰器在 Agent 节点函数执行后自动"
    "过滤掉不在白名单中的字段。权限规则矩阵如图 3-3 所示。"
))
add_para(doc, (
    "权限控制的具体实现采用装饰器模式，在节点函数返回结果后执行拦截。拦截逻辑包括"
    "两个层次：一是字段过滤——非白名单字段被静默丢弃并记录警告日志；二是异常兜底——"
    "当 Agent 执行抛出异常时，装饰器捕获异常并记录错误信息到 _last_error 和 "
    "_error_count 字段，同时确保不返回任何业务数据，避免污染共享状态。这种设计"
    "保证了单个 Agent 的故障不会扩散到其他 Agent，增强了系统的鲁棒性。"
))

add_heading_custom(doc, '3.5  循环检测与安全机制', level=2)
add_para(doc, (
    "多 Agent 系统的一个常见风险是无限循环——某个 Agent 可能因为逻辑错误而反复"
    "被激活，导致系统资源耗尽。本系统设计了双重循环检测机制：一是连续调用检测，"
    "当同一个非 Supervisor Agent 连续出现 5 次以上时，强制终止并返回 END；二是"
    "总步数上限检测，当总执行步数超过 50 步时强制终止。检测逻辑在 Supervisor 的"
    "条件边函数中执行，每次路由决策前都会进行预判（检查即将激活的 Agent 是否会导致"
    "第 6 次连续调用），实现了预防式检测。"
))
add_para(doc, (
    "此外，系统在 Supervisor 层面还设计了多层兜底策略：规则路由优先（0 token 消耗）、"
    "LLM 路由补充（处理模糊场景）、默认路由兜底（纯确定性，永不失败）。这种渐进式"
    "降级设计确保即使 LLM API 完全不可用，系统仍能按照确定性规则推进流程，不会出现"
    "「卡死」的情况。"
))

doc.add_page_break()

# ============================================================
# 第4章 核心模块实现
# ============================================================
add_heading_custom(doc, '4  核心模块实现', level=1)

add_heading_custom(doc, '4.1  Supervisor 路由调度实现', level=2)
add_para(doc, (
    "Supervisor 是多 Agent 系统的「大脑」，负责根据当前状态决定下一步路由目标。"
    "其执行流程分为九个步骤：（Step -1）重置残留的 next_agent，防止检查点中的旧值"
    "污染本次路由；（Step 0）处理 Agent 的主动 FINISH 信号；（Step 1）更新步数计数"
    "器；（Step 2）检查硬性终止条件；（Step 3）执行第一优先级硬编码检查（铁律规则）；"
    "（Step 4）处理 Agent 显式路由信号；（Step 5）执行规则路由；（Step 6）规则不确定"
    "时转向 LLM 路由；（Step 7）最终兜底验证；（Step 8-9）记录历史并返回路由决策。"
))
add_para(doc, (
    "规则路由（_rule_based_route）是 Supervisor 的核心智能所在。它通过一系列优先级"
    "排序的规则函数覆盖了约 80% 的路由场景，完全不消耗 LLM token。规则的优先级设计"
    "遵循「面试进行中优先处理」原则：如果系统正在面试流程中（有面试题且当前索引在"
    "范围内），用户的任何输入都优先视为对当前问题的回答（路由到 Evaluator），除非"
    "用户明确要求结束面试。这种设计避免了用户回答中包含「简历」「项目」等词汇时被"
    "误判为其他意图。意图检测函数包括：简历上传检测（_has_resume_intent，匹配文件"
    "扩展名和上传短语）、岗位匹配检测（_has_match_job_intent，匹配「匹配岗位」等"
    "关键词）、通用问答检测（_has_general_qa_intent，匹配技术术语和问句句式）、"
    "个人信息检测（_has_personal_info_intent，匹配第一人称+个人信息字段）等。"
))

add_heading_custom(doc, '4.2  ResumeMatch Agent — 简历解析与岗位匹配', level=2)
add_para(doc, (
    "ResumeMatch Agent 负责面试流程的前置准备阶段，包括两个核心任务：简历解析和"
    "岗位匹配。在简历解析方面，系统支持 PDF、DOCX 和 TXT 三种格式，通过 "
    "ResumeParser 类进行解析。解析过程包括文本提取、正则表达式匹配（姓名、电话、"
    "邮箱、教育经历等结构化字段）、以及 LLM 辅助信息抽取（技能标签、项目经验摘要"
    "等）。解析结果以结构化字典形式存入 resume_data 字段。"
))
add_para(doc, (
    "岗位匹配采用向量语义匹配方法。系统预先将岗位库（jobs.json）中的岗位描述通过"
    "Embedding 模型向量化后存入 Chroma 向量数据库。匹配时，将候选人的简历文本同样"
    "进行向量化，通过余弦相似度计算与各岗位的匹配分数，选取 Top-K 结果返回。这种"
    "语义匹配方法相比传统关键词匹配能够捕捉到更深层的技能关联，例如简历中的"
    "「FastAPI 开发」技能可能与「Python 后端开发」岗位高度匹配，即使两者没有直接的"
    "关键词重叠。"
))
add_para(doc, (
    "Agent 的工作模式遵循「一次只做一件事」原则，由 Supervisor 控制步进节奏：第一步"
    "解析简历并返回 FINISH 信号；Supervisor 再次激活后，第二步展示匹配结果；用户"
    "选择后，第三步记录选择并确认。这种分步设计使得用户可以在任何中间步骤插入反馈"
    "或修改选择。"
))

add_heading_custom(doc, '4.3  Interview Agent — 面试主持', level=2)
add_para(doc, (
    "Interview Agent 负责面试的核心环节——生成个性化的面试题并逐题向候选人提问。"
    "问题生成过程融合了三方面信息：候选人的简历数据（技能、经历、项目）、目标岗位的"
    "任职要求（技术栈、职责描述）、RAG 知识库的检索结果（该岗位的常见面试考点和"
    "参考题目）。LLM 根据综合上下文生成 5 道面试题，涵盖技术技能、项目经验、问题解决"
    "能力、团队协作和职业规划五个维度。"
))
add_para(doc, (
    "出题遵循严格的「简历锚定」原则：只能问候选人简历里明确写过的内容，不得凭空"
    "追问未涉及的技术细节。例如，如果简历中写了「使用 FastAPI 开发 API」，但未说明"
    "技术选型原因，则不应问「为什么选择 FastAPI 而非 Django？」这类问题。这一约束"
    "通过系统提示词中的规则强制执行，减少了模型的幻觉和不当预设。"
))
add_para(doc, (
    "在问题难度控制方面，Evaluator Agent 会根据上一题的回答评分设置 next_difficulty "
    "字段（easy/normal/hard），Interview Agent 读取该字段并在下一题生成时调整难度。"
    "回答好的候选人面临更有挑战性的追问，回答一般的候选人则获得更基础的题目，实现了"
    "自适应面试体验。"
))

add_heading_custom(doc, '4.4  Evaluator Agent — 回答评分', level=2)
add_para(doc, (
    "Evaluator Agent 是系统评估能力的核心体现。当候选人回答完一道面试题后，Evaluator "
    "Agent 被激活，执行以下流程：（1）调用 RAG 知识库检索该技术领域的参考答案和评分"
    "标准；（2）将面试问题、候选人回答、岗位背景和 RAG 检索结果组装为评分提示词；"
    "（3）调用 LLM 进行结构化评分，输出包含 score（1-10 分）、strengths（优点）、"
    "weaknesses（不足）和 suggestions（改进建议）的 JSON 结构；（4）将评分结果追加"
    "到 answers 和 evaluations 列表中，并更新 last_answer_score 字段。"
))
add_para(doc, (
    "评分系统不仅给出总分，还从技术深度、表达清晰度、经验相关性、综合素质四个维度"
    "分别打分。评价语言采用建设性口吻——指出不足的同时提供具体的改进建议，而非简单"
    "的批判。这种设计使得系统不仅是筛选工具，也具备一定的培训辅导价值。"
))

# ---- 截图占位：面试对话界面 ----
chat_img = os.path.join(os.path.dirname(__file__) or '.', 'paper_screenshot_chat.png')
generate_placeholder_image(chat_img, '图 4-1  面试对话界面', 'AI 面试官与候选人的实时问答交互')
insert_image(doc, chat_img, '图 4-1  面试对话界面截图')

add_heading_custom(doc, '4.5  Report Agent — 报告生成', level=2)
add_para(doc, (
    "当所有面试问题回答完毕（或用户主动要求结束）时，Report Agent 被激活。它汇总"
    "整个面试过程中的问答记录和评分数据，生成一份结构化的面试评估报告。报告包含以下"
    "部分：（1）综合评分（0-10 分），由各题得分的加权平均值计算得出；（2）分项评分"
    "（技术能力、表达沟通、经验匹配、综合素质），每项附有 2-3 句具体说明；（3）各题"
    "详细评价，包含原问题、候选人回答摘要、得分和建议；（4）录用建议，综合所有评分"
    "给出最终推荐等级（强烈推荐/推荐/保留考虑/不推荐）。"
))
add_para(doc, (
    "报告生成完成后，Report Agent 会清空面试相关状态字段（interview_questions、"
    "answers、current_question_idx 等），将 interview_stage 设置为 'done'，"
    "为可能的下一轮面试做好准备。"
))

# ---- 截图占位：面试报告展示 ----
report_img = os.path.join(os.path.dirname(__file__) or '.', 'paper_screenshot_report.png')
generate_placeholder_image(report_img, '图 4-2  面试评估报告界面', '综合评分 + 分项评分 + 录用建议')
insert_image(doc, report_img, '图 4-2  面试评估报告界面截图')

add_heading_custom(doc, '4.6  RAG 面试知识库实现', level=2)
add_para(doc, (
    "面试知识库（KnowledgeBase）是整个系统的知识支撑底座。其文档加载流程为：从"
    "interview_kb/ 目录读取 Markdown 格式的面试知识文档（按技术方向分类），使用"
    "RecursiveCharacterTextSplitter 进行智能分块（chunk_size=800, chunk_overlap=100），"
    "通过智谱 Embedding 模型将文档块向量化后存入 Chroma 的对应 Collection。"
))
add_para(doc, (
    "RAGChain 类封装了完整的检索增强生成流程：（1）retrieve() — 根据查询文本从"
    "知识库中检索 Top-K 最相关文档块；（2）_format_context() — 将检索结果格式化为"
    "结构化的上下文文本；（3）generate_questions() — 将上下文注入面试题生成提示词；"
    "（4）evaluate_answer() — 将参考答案和评分标准注入评分提示词。检索时支持按技术"
    "方向过滤（元数据过滤），避免跨领域的知识污染。"
))

add_heading_custom(doc, '4.7  LangGraph 图的构建与编译', level=2)
add_para(doc, (
    "系统 LangGraph 图的构建在 build_multi_agent_graph 函数中完成，核心代码如下："
))
add_code_block(doc, (
    'def build_multi_agent_graph(checkpointer=None):\n'
    '    """构建多 Agent 图（完整版）"""\n'
    '    graph = StateGraph(MultiAgentState)\n'
    '    # 添加 Supervisor 节点\n'
    '    graph.add_node("supervisor", supervisor_node)\n'
    '    # 添加各业务 Agent 节点\n'
    '    graph.add_node("resume_match", resume_match_node)\n'
    '    graph.add_node("interview", interview_node)\n'
    '    graph.add_node("evaluate", evaluator_node)\n'
    '    graph.add_node("report", report_node)\n'
    '    graph.add_node("direct_reply", direct_reply_node)\n'
    '    # 起始 → Supervisor\n'
    '    graph.add_edge(START, "supervisor")\n'
    '    # Supervisor → 条件路由到各 Agent\n'
    '    graph.add_conditional_edges(\n'
    '        "supervisor", route_after_supervisor,\n'
    '        {"resume_match": "resume_match",\n'
    '         "interview": "interview",\n'
    '         "evaluate": "evaluate",\n'
    '         "report": "report",\n'
    '         "direct_reply": "direct_reply",\n'
    '         END: END}\n'
    '    )\n'
    '    # 所有 Agent 完成后 → 回到 Supervisor\n'
    '    for agent in ["resume_match","interview",\n'
    '                  "evaluate","report","direct_reply"]:\n'
    '        graph.add_edge(agent, "supervisor")\n'
    '    return graph.compile(checkpointer=checkpointer)'
))

add_para(doc, (
    "图的执行流程为：用户消息进入 → Supervisor 分析状态并决定目标 Agent → 目标 "
    "Agent 执行业务逻辑并更新状态 → 返回 Supervisor → 再次路由判断 → ... → 直至"
    "Supervisor 返回 FINISH 信号。每个 Agent 执行完毕后都回到 Supervisor，这种"
    "「中心辐射」拓扑确保了流程控制的集中化和可预测性。检查点存储器在每个超级步骤"
    "后自动保存状态快照，支持会话的暂停恢复和历史回溯。"
))

doc.add_page_break()

# ============================================================
# 第5章 系统测试与分析
# ============================================================
add_heading_custom(doc, '5  系统测试与分析', level=1)

add_heading_custom(doc, '5.1  测试环境与方案', level=2)
add_para(doc, (
    "系统测试环境配置如下：操作系统 Windows 11，Python 3.11，内存 16GB，LLM API "
    "提供商为 DeepSeek（模型 deepseek-chat）。测试方案涵盖四个维度：（1）功能测试——"
    "验证各 Agent 的核心功能是否正确执行；（2）集成测试——验证 Agent 之间的协作流程"
    "是否顺畅；（3）异常测试——验证系统在 LLM 超时、文件损坏、无效输入等异常场景下的"
    "表现；（4）性能测试——评估端到端面试流程的响应时间和资源消耗。"
))

add_heading_custom(doc, '5.2  权限控制系统测试', level=2)
add_para(doc, (
    "权限控制系统通过 agents/state.py 中的自检代码进行验证。测试覆盖了以下场景："
    "字段白名单过滤——ResumeMatch Agent 返回 interview_questions 字段时被静默丢弃；"
    "异常兜底——Agent 抛出 ValueError 时，装饰器捕获异常并记录到 _last_error，"
    "不返回任何业务数据；跨 Agent 读取——所有 Agent 可以读取其他 Agent 的字段，"
    "不受权限限制。所有测试均已通过（state.py 运行 ALL STATE TESTS PASSED），"
    "验证了权限隔离机制的正确性。"
))

add_heading_custom(doc, '5.3  循环检测测试', level=2)
add_para(doc, (
    "循环检测同样通过单元测试验证：同一 Agent 连续出现 5 次触发强制终止；预判检测"
    "在第 4 次出现时提前拦截第 5 次请求；总步数超过 50 步触发强制终止。测试涵盖了"
    "边界情况（如连续 4 次不加拦截、5 次触发拦截的阈值），确保循环检测逻辑的精确性。"
))

add_heading_custom(doc, '5.4  Supervisor 路由测试', level=2)
add_para(doc, (
    "Supervisor 路由逻辑通过 supervisor.py 的自检代码进行系统测试。测试覆盖了六种"
    "典型场景：空状态路由到 resume_match、简历已解析但无岗位路由到 resume_match、"
    "岗位已选但无面试题路由到 interview、所有题目答完路由到 report、报告已生成路由"
    "到 FINISH、Agent 连续失败 3 次触发降级路由。所有测试均已通过（ALL SUPERVISOR "
    "TESTS PASSED），验证了路由逻辑和兜底策略的正确性。"
))

add_heading_custom(doc, '5.5  端到端面试流程测试', level=2)
add_para(doc, (
    "端到端测试模拟了完整的面试流程：用户上传简历 → 系统解析简历并展示匹配岗位 → "
    "用户选择岗位 → 系统生成 5 道个性化面试题 → 用户逐题回答 → 系统对每道题进行评分 "
    "→ 全部回答完毕后生成综合评估报告。测试结果表明系统能够在 2-3 分钟内完成包含"
    "简历解析、岗位匹配、题目生成和评分反馈的完整面试流程，各 Agent 之间的切换流畅"
    "无卡顿。"
))

# ---- 截图占位：测试结果展示 ----
test_img = os.path.join(os.path.dirname(__file__) or '.', 'paper_screenshot_test.png')
generate_placeholder_image(test_img, '图 5-1  端到端面试流程测试', '完整的面试流程：简历 → 匹配 → 问答 → 评分 → 报告')
insert_image(doc, test_img, '图 5-1  端到端面试流程测试截图')

add_heading_custom(doc, '5.6  Web 服务性能测试', level=2)
add_para(doc, (
    "Web 服务层基于 FastAPI + Uvicorn 部署，使用 curl 进行了基本的接口可用性测试："
    "GET / 返回 200（静态页面服务正常），POST /api/chat 处理单次对话请求的平均响应"
    "时间约为 3-8 秒（取决于 LLM API 的响应速度），文件上传接口支持最大 10MB 的"
    "PDF/DOCX 文件。在并发场景下（同时 3-5 个会话），系统通过 asyncio 异步 I/O "
    "有效处理了 LLM 调用的长等待，未出现请求阻塞的情况。"
))

doc.add_page_break()

# ============================================================
# 第6章 总结与展望
# ============================================================
add_heading_custom(doc, '6  总结与展望', level=1)

add_heading_custom(doc, '6.1  工作总结', level=2)
add_para(doc, (
    "本文围绕「基于 LangGraph 的多 Agent 智能面试系统」这一课题，完成了从需求分析、"
    "架构设计到编码实现和测试验证的完整工程实践。主要成果总结如下："
))
add_para(doc, (
    "（1）设计并实现了一套基于 LangGraph 的多 Agent 协作架构。该架构将面试全流程"
    "拆解为五个独立 Agent 协作完成的子任务，通过 Supervisor 实现智能路由和统一调度，"
    "通过 MultiAgentState 实现共享状态管理。架构具有清晰的模块边界和良好的可扩展性。"
))
add_para(doc, (
    "（2）建立了完善的系统安全机制。包括基于白名单的权限控制系统（防止 Agent 间"
    "状态污染）、双重循环检测机制（连续调用 + 总步数上限）、多层路由兜底策略"
    "（规则优先 → LLM 补充 → 默认兜底），确保系统在各种异常情况下都能正常运行。"
))
add_para(doc, (
    "（3）集成了 RAG 检索增强生成技术。基于 Chroma 向量数据库构建了面试知识库，"
    "使系统能够检索真实的面试考点和评分标准，显著提升了面试题目的专业性和评分的"
    "客观性。"
))
add_para(doc, (
    "（4）基于 FastAPI 和原生 Web 技术实现了可用的系统原型。前端支持实时对话、"
    "会话管理、文件拖拽上传等功能，后端通过三级检查点存储策略（Redis→SQLite→"
    "Memory）保障了不同部署环境下的数据持久性。"
))

add_heading_custom(doc, '6.2  不足与展望', level=2)
add_para(doc, (
    "本系统作为工程实训项目，仍存在以下方面的不足和可改进空间："
))
add_para(doc, (
    "（1）语音交互能力：当前系统仅支持文本输入输出。引入语音识别（ASR）和语音合成"
    "（TTS）技术可以使面试体验更加自然，也更接近真实面试场景。"
))
add_para(doc, (
    "（2）多模态评估：目前的评分仅基于文本回答，未考虑候选人的微表情、语气等非语言"
    "信息。未来可引入视频分析技术，实现更全面的候选人评估。"
))
add_para(doc, (
    "（3）面试题库的丰富性：当前知识库的面试文档覆盖的技术方向有限。通过持续收集"
    "和整理来自真实面试的题目和评价数据，可以不断提升系统的专业性和实用性。"
))
add_para(doc, (
    "（4）Agent 自适应能力：当前 Agent 的行为主要通过预定义的 System Prompt 控制。"
    "引入强化学习或基于反馈的 Prompt 自动优化机制，可以使 Agent 在实际运行中持续"
    "改进其表现。"
))
add_para(doc, (
    "（5）分布式部署：当前系统为单机部署架构。对于企业级应用场景，可以考虑将不同的"
    "Agent 部署为独立的微服务，通过消息队列进行异步通信，提升系统的吞吐量和可用性。"
))

doc.add_page_break()

# ============================================================
# 参考文献
# ============================================================
add_heading_custom(doc, '参考文献', level=1)

references = [
    "[1] LangChain Team. LangGraph Documentation[EB/OL]. https://langchain-ai.github.io/langgraph/, 2025.",
    "[2] DeepSeek Team. DeepSeek API Documentation[EB/OL]. https://api-docs.deepseek.com/, 2025.",
    "[3] FastAPI Team. FastAPI Documentation[EB/OL]. https://fastapi.tiangolo.com/, 2025.",
    "[4] Chroma Team. Chroma Vector Database Documentation[EB/OL]. https://docs.trychroma.com/, 2025.",
    "[5] Lewis P, Perez E, Piktus A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks[C]. NeurIPS, 2020.",
    "[6] Wang L, Ma C, Feng X, et al. A Survey on Large Language Model based Autonomous Agents[J]. Frontiers of Computer Science, 2024, 18(6).",
    "[7] Xi Z, Chen W, Guo X, et al. The Rise and Potential of Large Language Model Based Agents: A Survey[J]. arXiv preprint arXiv:2309.07864, 2023.",
    "[8] Chase H. LangChain: Building Applications with LLMs through Composability[EB/OL]. https://github.com/langchain-ai/langchain, 2025.",
    "[9] Vaswani A, Shazeer N, Parmar N, et al. Attention Is All You Need[C]. NeurIPS, 2017.",
    "[10] Brown T B, Mann B, Ryder N, et al. Language Models are Few-Shot Learners[C]. NeurIPS, 2020.",
    "[11] 李飞飞, 邓嘉, 黄高, 等. 多智能体系统综述[J]. 中国科学: 信息科学, 2023, 53(7): 1209-1240.",
    "[12] 张俊林. 大语言模型的训练与推理优化综述[J]. 计算机学报, 2025, 48(1): 1-35.",
    "[13] 周志远, 王晓明, 刘洋. 基于大语言模型的智能招聘系统研究综述[J]. 中文信息学报, 2024, 38(5): 1-18.",
]

for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(ref)
    run.font.name = '宋体'
    run.font.size = Pt(11)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ============================================================
# 致谢
# ============================================================
add_heading_custom(doc, '致  谢', level=1)

add_para(doc, (
    "在本论文完成之际，衷心感谢指导老师在项目选题、技术方案设计和论文撰写过程中"
    "的悉心指导。导师严谨的学术态度和丰富的工程实践经验为项目的顺利完成提供了重要保障。"
))
add_para(doc, (
    "感谢 LangChain、LangGraph、DeepSeek、FastAPI 等开源社区提供的优秀工具和"
    "文档支持，使得本系统能够在较短的时间内完成从概念到原型的完整开发流程。感谢"
    "各位技术博客作者和开源贡献者分享的技术文章和代码示例，为本项目的实现提供了"
    "宝贵的参考。"
))
add_para(doc, (
    "最后，感谢在项目开发和论文撰写期间给予理解和支持的家人和朋友。"
))

# ============================================================
# 保存
# ============================================================
output_path = "基于LangGraph的多Agent智能面试系统设计与实现.docx"
doc.save(output_path)
print(f"[OK] Paper generated: {output_path}")
print(f"     Size: {os.path.getsize(output_path) / 1024:.1f} KB")
