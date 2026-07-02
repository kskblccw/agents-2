#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
简历解析器 (resume_parser.py)
============================
从 PDF/Word 文件中提取简历信息，并转换为系统需要的 JSON 格式

功能：
1. 支持 PDF 文件读取（使用 pdfplumber）
2. 支持 Word 文件读取（使用 python-docx）
3. 使用 LLM（GLM-4）进行信息提取
4. 输出标准化的 JSON 格式

使用方法：
    parser = ResumeParser()
    result = parser.parse_file('resume.pdf')
    # 或
    result = parser.parse_text(raw_text)
"""

import os
import re
import json
from typing import Optional, Dict, Any, List

# =============================================================================
# 第一部分：文件读取函数
# =============================================================================

def read_pdf(file_path: str) -> str:
    """
    从 PDF 文件中提取文本
    
    参数：
        file_path: PDF 文件路径
    
    返回：
        str: 提取的文本内容
    
    异常：
        ImportError: 如果 pdfplumber 未安装
        FileNotFoundError: 如果文件不存在
    """
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("请安装 pdfplumber: pip install pdfplumber")
    
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")
    
    text_content = []
    
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_content.append(page_text)
    
    return "\n".join(text_content)


def read_docx(file_path: str) -> str:
    """
    从 Word 文件中提取文本
    
    参数：
        file_path: Word 文件路径
    
    返回：
        str: 提取的文本内容
    
    异常：
        ImportError: 如果 python-docx 未安装
        FileNotFoundError: 如果文件不存在
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("请安装 python-docx: pip install python-docx")
    
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")
    
    doc = Document(file_path)
    paragraphs = []
    
    # 提取段落文本
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            paragraphs.append(text)
    
    # 尝试提取表格内容
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
            if row_text:
                paragraphs.append(row_text)
    
    # 尝试提取页眉和页脚
    for section in doc.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                text = paragraph.text.strip()
                if text and text not in paragraphs:
                    paragraphs.append(text)
        if section.footer:
            for paragraph in section.footer.paragraphs:
                text = paragraph.text.strip()
                if text and text not in paragraphs:
                    paragraphs.append(text)
    
    # 调试信息：打印提取的段落数量
    print(f"从 DOCX 文件中提取了 {len(paragraphs)} 个段落")
    
    result = "\n".join(paragraphs)
    
    # 如果提取的文本过少，尝试其他方法
    if len(result.strip()) < 20:
        print(f"警告：DOCX 文件内容较少，尝试从 XML 直接提取...")
        result = _extract_docx_text_from_xml(file_path)
    
    return result


def _extract_docx_text_from_xml(file_path: str) -> str:
    """
    从 DOCX 文件的 XML 结构中直接提取文本（备用方法）
    
    当 python-docx 无法正确读取时使用此方法
    """
    import zipfile
    from xml.etree import ElementTree as ET
    
    namespace = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    
    try:
        with zipfile.ZipFile(file_path, 'r') as docx:
            # 读取主文档内容
            with docx.open('word/document.xml') as f:
                tree = ET.parse(f)
            
            root = tree.getroot()
            paragraphs = []
            
            # 提取所有文本段落
            for elem in root.iter(namespace + 't'):
                if elem.text and elem.text.strip():
                    paragraphs.append(elem.text.strip())
            
            # 读取页眉
            try:
                with docx.open('word/header1.xml') as f:
                    tree = ET.parse(f)
                for elem in tree.getroot().iter(namespace + 't'):
                    if elem.text and elem.text.strip():
                        paragraphs.append(elem.text.strip())
            except KeyError:
                pass  # 可能没有页眉
            
            # 读取页脚
            try:
                with docx.open('word/footer1.xml') as f:
                    tree = ET.parse(f)
                for elem in tree.getroot().iter(namespace + 't'):
                    if elem.text and elem.text.strip():
                        paragraphs.append(elem.text.strip())
            except KeyError:
                pass  # 可能没有页脚
            
            print(f"从 XML 中提取了 {len(paragraphs)} 个文本片段")
            return "\n".join(paragraphs)
    
    except Exception as e:
        print(f"从 XML 提取文本失败: {str(e)}")
        return ""


def extract_text_from_file(file_path: str) -> str:
    """
    根据文件扩展名自动选择合适的读取方法
    
    参数：
        file_path: 文件路径
    
    返回：
        str: 提取的文本内容
    """
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.pdf':
        return read_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        return read_docx(file_path)
    elif ext == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    else:
        raise ValueError(f"不支持的文件格式: {ext}。支持的格式: PDF, DOCX, DOC, TXT")


# =============================================================================
# 第二部分：LLM 信息提取
# =============================================================================

def ask_llm(prompt: str, model: str = "deepseek-v4-pro") -> str:
    """
    调用阿里云百炼 LLM API（OpenAI 兼容模式）

    参数：
        prompt: 提示词
        model: 模型名称，默认 deepseek-chat

    返回：
        str: API 返回的回答
    """
    from openai import OpenAI

    api_key = os.getenv("DEEPSEEK_API_KEY")
    if not api_key:
        raise ValueError("未设置 DEEPSEEK_API_KEY 环境变量")

    client = OpenAI(
        api_key=api_key,
        base_url="https://api.deepseek.com/v1",
        timeout=120.0,
    )

    try:
        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,  # 简历解析用低温度，保证输出稳定
        )
        return response.choices[0].message.content
    except Exception as e:
        raise RuntimeError(f"LLM API 调用失败: {str(e)}")


# 保留旧名称作为别名，兼容现有调用
def ask_glm(prompt: str, model: str = "deepseek-chat") -> str:
    """兼容旧接口，内部调用 ask_llm"""
    return ask_llm(prompt, model=model)


def parse_resume_with_llm(text: str) -> Dict[str, Any]:
    """
    使用 LLM 从简历文本中提取结构化信息
    
    参数：
        text: 简历文本内容
    
    返回：
        dict: 结构化的简历信息，包含以下字段：
            - name: 姓名
            - skills: 技能列表
            - experience: 工作经历列表
            - education: 教育背景
            - contact: 联系方式
            - summary: 个人简介
            - raw_text: 原始文本（保留）
    """
    prompt = f"""请从以下简历文本中提取结构化信息，并以 JSON 格式输出。

【简历文本】
{text}

请返回以下 JSON 格式（只输出 JSON，不要有其他文字）：
{{
    "name": "姓名（如果未找到，填写空字符串）",
    "skills": ["技能1", "技能2", "技能3"],
    "experience": [
        {{
            "company": "公司名称",
            "position": "职位名称", 
            "duration": "工作时间（如：2020.01-2023.01）",
            "description": "工作描述"
        }}
    ],
    "education": [
        {{
            "school": "学校名称",
            "degree": "学历（如：本科、硕士）",
            "major": "专业",
            "graduation": "毕业时间"
        }}
    ],
    "contact": {{
        "phone": "电话号码",
        "email": "邮箱",
        "location": "所在地"
    }},
    "summary": "个人简介（1-3句话）"
}}

注意：
1. 如果某项信息未找到，填写空字符串或空数组
2. skills 必须是字符串数组
3. experience 和 education 必须是数组
4. 只输出 JSON，不要有任何其他文字
"""

    try:
        response = ask_glm(prompt)
        
        # 尝试解析 JSON
        json_str = response.strip()
        
        # 处理可能存在的 markdown 代码块
        if json_str.startswith('```'):
            json_str = json_str.split('```')[1]
            if json_str.startswith('json'):
                json_str = json_str[4:]
            json_str = json_str.strip()
        
        # 找到 JSON 对象的开始和结束
        start = json_str.find('{')
        end = json_str.rfind('}')
        if start != -1 and end != -1:
            json_str = json_str[start:end+1]
        
        result = json.loads(json_str)
        
        # 验证必要字段并添加默认值
        structured = {
            'name': result.get('name', ''),
            'skills': result.get('skills', []),
            'experience': result.get('experience', []),
            'education': result.get('education', []),
            'contact': result.get('contact', {}),
            'summary': result.get('summary', ''),
            'text': text,  # 保留原始文本用于职位匹配
            'raw_text': text  # 保留原始文本（兼容旧代码）
        }
        
        return structured
        
    except json.JSONDecodeError as e:
        # JSON 解析失败，打印详细错误信息方便调试
        print(f"\n{'='*60}")
        print(f"警告：JSON 解析失败")
        print(f"错误类型: JSONDecodeError")
        print(f"错误信息: {str(e)}")
        print(f"LLM 返回的原始内容长度: {len(response) if 'response' in dir() else 'N/A'}")
        if 'response' in dir() and response:
            # 截取前500个字符显示
            preview = response[:500] if len(response) > 500 else response
            print(f"LLM 返回内容预览:\n{preview}")
        print(f"{'='*60}\n")
        # 返回降级结果
        return fallback_parse(text)
    except Exception as e:
        # 其他异常，打印详细错误信息
        print(f"\n{'='*60}")
        print(f"警告：LLM 调用或解析失败")
        print(f"错误类型: {type(e).__name__}")
        print(f"错误信息: {str(e)}")
        print(f"{'='*60}\n")
        return fallback_parse(text)


def fallback_parse(text: str) -> Dict[str, Any]:
    """
    备用解析方法：当 LLM 不可用时使用正则表达式提取信息
    
    参数：
        text: 简历文本内容
    
    返回：
        dict: 结构化的简历信息
    """
    result = {
        'name': '',
        'skills': [],
        'experience': [],
        'education': [],
        'contact': {},
        'summary': '',
        'text': text,  # 保留原始文本用于职位匹配
        'raw_text': text
    }
    
    # 提取姓名（通常在文档开头，较大的字号后跟随的文字）
    name_pattern = r'^[\u4e00-\u9fa5]{2,4}(?=\s|$|，|。|：|:)'
    name_match = re.search(name_pattern, text[:100])
    if name_match:
        result['name'] = name_match.group(0)
    
    # 提取技能（查找常见技能关键词）
    skill_keywords = [
        'Python', 'Java', 'JavaScript', 'C++', 'C#', 'Go', 'Rust', 'PHP', 'Ruby', 'Swift',
        'HTML', 'CSS', 'React', 'Vue', 'Angular', 'Node.js', 'Django', 'Flask', 'Spring',
        'MySQL', 'PostgreSQL', 'MongoDB', 'Redis', 'Elasticsearch',
        'Git', 'Docker', 'Kubernetes', 'Jenkins', 'Linux', 'AWS', 'Azure', 'GCP',
        '机器学习', '深度学习', 'AI', 'NLP', '计算机视觉', '数据分析', '大数据',
        '微服务', '分布式', '高并发', '缓存', '消息队列'
    ]
    
    found_skills = []
    text_lower = text.lower()
    for skill in skill_keywords:
        if skill.lower() in text_lower:
            found_skills.append(skill)
    
    result['skills'] = found_skills[:10]  # 最多保留10个技能
    
    # 提取电话
    phone_pattern = r'1[3-9]\d{9}'
    phone_match = re.search(phone_pattern, text)
    if phone_match:
        result['contact']['phone'] = phone_match.group(0)
    
    # 提取邮箱
    email_pattern = r'[\w.-]+@[\w.-]+\.\w+'
    email_match = re.search(email_pattern, text)
    if email_match:
        result['contact']['email'] = email_match.group(0)
    
    return result


# =============================================================================
# 第三部分：简历解析器主类
# =============================================================================

class ResumeParser:
    """
    简历解析器类
    
    支持从 PDF/Word 文件或文本中提取简历信息，
    并转换为系统需要的 JSON 格式。
    """
    
    def __init__(self):
        """初始化简历解析器"""
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt']
    
    def parse_file(self, file_path: str) -> Dict[str, Any]:
        """
        解析简历文件
        
        参数：
            file_path: 简历文件路径（支持 PDF、DOCX、TXT）
        
        返回：
            dict: 结构化的简历信息
        
        异常：
            FileNotFoundError: 如果文件不存在
            ValueError: 如果文件格式不支持
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        ext = os.path.splitext(file_path)[1].lower()
        if ext not in self.supported_formats:
            raise ValueError(
                f"不支持的文件格式: {ext}。"
                f"支持的格式: {', '.join(self.supported_formats)}"
            )
        
        # 提取文本
        print(f"正在读取文件: {file_path}")
        text = extract_text_from_file(file_path)
        
        # 调试：打印提取的文本长度
        print(f"提取的文本长度: {len(text)} 字符")
        
        if not text or len(text.strip()) < 20:
            raise ValueError("文件内容过少或无法提取文本")
        
        # 调试：显示文本预览（前200字符）
        if len(text) > 200:
            preview = text[:200] + "..."
        else:
            preview = text
        print(f"文本内容预览:\n{preview}\n")
        
        # 使用 LLM 提取信息
        print("正在使用 AI 提取简历信息...")
        result = parse_resume_with_llm(text)
        
        # 添加元数据
        result['meta'] = {
            'source_file': os.path.basename(file_path),
            'file_type': ext,
            'text_length': len(text)
        }
        
        return result
    
    def parse_text(self, text: str) -> Dict[str, Any]:
        """
        解析简历文本
        
        参数：
            text: 简历文本内容
        
        返回：
            dict: 结构化的简历信息
        """
        if not text or len(text.strip()) < 20:
            raise ValueError("文本内容过少")
        
        # 使用 LLM 提取信息
        print("正在使用 AI 提取简历信息...")
        result = parse_resume_with_llm(text)
        
        # 添加元数据
        result['meta'] = {
            'source_file': None,
            'file_type': 'text',
            'text_length': len(text)
        }
        
        return result
    
    def save_to_json(self, data: Dict[str, Any], output_path: str) -> None:
        """
        将解析结果保存为 JSON 文件
        
        参数：
            data: 解析结果
            output_path: 输出文件路径
        """
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        
        print(f"结果已保存到: {output_path}")
    
    def get_skills_text(self, data: Dict[str, Any]) -> str:
        """
        获取格式化的技能文本（用于向量化）
        
        参数：
            data: 解析结果
        
        返回：
            str: 格式化的技能文本
        """
        skills = data.get('skills', [])
        if isinstance(skills, list):
            skills_text = ', '.join(skills)
        else:
            skills_text = str(skills)
        
        return skills_text
    
    def get_experience_text(self, data: Dict[str, Any]) -> str:
        """
        获取格式化的经历文本（用于向量化）
        
        参数：
            data: 解析结果
        
        返回：
            str: 格式化的经历文本
        """
        experience = data.get('experience', [])
        if not isinstance(experience, list):
            return str(experience)
        
        experience_parts = []
        for exp in experience:
            if isinstance(exp, dict):
                part = f"{exp.get('company', '')} {exp.get('position', '')} {exp.get('duration', '')}"
                desc = exp.get('description', '')
                if desc:
                    part += f": {desc}"
                experience_parts.append(part)
            else:
                experience_parts.append(str(exp))
        
        return '\n'.join(experience_parts)


# =============================================================================
# 第四部分：示例用法和测试
# =============================================================================

def demo():
    """
    演示简历解析器的使用方法
    """
    parser = ResumeParser()
    
    # 示例1：从文本解析
    sample_text = """
    张三
    电话: 13800138000
    邮箱: zhangsan@example.com
    
    个人简介:
    5年Python开发经验，熟悉Web开发和数据分析。
    
    技能:
    Python, Django, Flask, MySQL, Redis, Docker, Git, Linux
    
    工作经历:
    2020.01 - 至今 字节跳动 高级Python开发工程师
    负责后端系统开发，使用Django和Flask构建高性能Web服务。
    
    2018.01 - 2020.01 阿里巴巴 Python开发工程师
    参与电商平台后端开发，使用Python和MySQL。
    
    教育背景:
    2014.09 - 2018.06 北京大学 计算机科学与技术 本科
    """
    
    print("=" * 70)
    print("简历解析器演示")
    print("=" * 70)
    
    try:
        result = parser.parse_text(sample_text)
        
        print("\n解析结果：")
        print(f"姓名: {result['name']}")
        print(f"技能: {', '.join(result['skills'])}")
        print(f"工作经历数: {len(result['experience'])}")
        print(f"教育背景数: {len(result['education'])}")
        print(f"联系方式: {result['contact']}")
        
        # 保存为 JSON 文件
        parser.save_to_json(result, 'parsed_resume.json')
        
    except Exception as e:
        print(f"演示失败: {str(e)}")
        print("\n提示：请确保设置了 ZHIPUAI_API_KEY 环境变量")


if __name__ == "__main__":
    demo()
